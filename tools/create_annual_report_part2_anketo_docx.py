r"""
年報 第2部（アンケート編）Word文書生成スクリプト

コンテンツ源（Markdownファイル）を読み込み、Word文書として出力します。
図（グラフ）は Excelで作成済みのものを手動で貼り付ける前提のため、
[CHART:xxx] プレースホルダーは「▼ここに図を貼り付け▼」枠として出力されます。

【AI assisted 自治会運営の設計思想】
  テキストを修正したい  → annual_report_part2_anketo_draft.md を編集してこのスクリプトを実行
  来年度に引き継ぐ場合  → 同じ手順で再生成できる

Requirements:
    python-docx, openpyxl（いずれもインストール済み）

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_part2_anketo_docx.py

データソース:
    docs/community/2026__/annual_report_part2_anketo_draft.md  ← テキスト

出力:
    docs/community/2026__/annual_report_part2_anketo.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# パス設定
# ============================================================
BASE = Path(__file__).parent.parent
DRAFT_FILE  = BASE / 'docs' / 'community' / '2026__' / 'annual_report_part2_anketo_draft.md'
OUTPUT_FILE = BASE / 'docs' / 'community' / '2026__' / 'annual_report_part2_anketo.docx'



# ============================================================
# Markdownパーサー
# ============================================================
def parse_draft(md_path):
    """
    Markdownの「## ドラフト本文」セクションを解析し、
    段落・見出し・チャートプレースホルダーのリストを返す。

    返り値の要素:
        {'type': 'h2',    'content': '...'}
        {'type': 'h3',    'content': '...'}
        {'type': 'h4',    'content': '...'}
        {'type': 'para',  'content': '...'}  ← 複数行は改行で結合
        {'type': 'chart', 'key': '夏_空き地利用'}
    """
    text = md_path.read_text(encoding='utf-8')

    # 「## ドラフト本文」〜「## 執筆メモ」を抽出
    m = re.search(r'## ドラフト本文\n+(.*?)(?=\n## 執筆メモ|\Z)', text, re.DOTALL)
    if not m:
        raise ValueError("「## ドラフト本文」セクションが見つかりません")

    lines = m.group(1).split('\n')
    items = []
    buf = []  # 引用ブロックのバッファ

    def flush_buf():
        if buf:
            # バッファ内の行を結合（>区切りの段落ごとに保存済み）
            items.append({'type': 'para', 'content': '\n'.join(buf)})
            buf.clear()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == '---':
            flush_buf()
        elif stripped.startswith('#### '):
            flush_buf()
            items.append({'type': 'h4', 'content': stripped[5:].strip()})
        elif stripped.startswith('### '):
            flush_buf()
            items.append({'type': 'h3', 'content': stripped[4:].strip()})
        elif stripped.startswith('## '):
            flush_buf()
            items.append({'type': 'h2', 'content': stripped[3:].strip()})
        elif stripped.startswith('[CHART:') and stripped.endswith(']'):
            flush_buf()
            key = stripped[7:-1]
            items.append({'type': 'chart', 'key': key})
        elif stripped.startswith('>'):
            text_part = stripped[1:].strip()
            if text_part:
                buf.append(text_part)
            else:
                # 空の > は段落区切り
                flush_buf()
        elif stripped:
            buf.append(stripped)
        else:
            flush_buf()

        i += 1

    flush_buf()
    return items

# ============================================================
# インラインMarkdown（太字のみ）パーサー
# ============================================================
def parse_inline(text):
    """**bold** テキストを [(text, is_bold), ...] に変換"""
    result = []
    pos = 0
    for m in re.finditer(r'\*\*(.*?)\*\*', text):
        if m.start() > pos:
            result.append((text[pos:m.start()], False))
        result.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        result.append((text[pos:], False))
    return result if result else [(text, False)]

# ============================================================
# Word文書生成
# ============================================================
def set_default_font(doc, font_name='游明朝', size_pt=11):
    """文書デフォルトフォントを設定"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    # 東アジア文字フォント設定
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_para(doc, text, font_size=11, bold=False, indent=False, space_before=None, space_after=6):
    """段落を追加（インライン太字対応）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)

    for segment_text, is_bold in parse_inline(text):
        run = p.add_run(segment_text)
        run.bold = is_bold or bold
        run.font.size = Pt(font_size)

    return p

def add_heading(doc, text, level, font_size):
    """見出しを追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 3 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if level == 3:
        run.font.color.rgb = RGBColor(0x17, 0x37, 0x5E)  # 濃紺
    elif level == 4:
        run.font.color.rgb = RGBColor(0x26, 0x4F, 0x78)  # 中濃紺
    return p

def add_chart_placeholder(doc, key):
    """図の手動貼り付け枠を挿入する"""
    LABELS = {
        '夏_空き地利用': '図：空き地利用への意向（夏アンケート）',
        '夏_地域要望':   '図：地域要望の集計（夏アンケート・複数回答可）',
        '冬_会費改定':   '図：会費改定案への意向（冬アンケート・3択）',
        '冬_参考項目':   '図：参考項目の集計（役員会委任・前提支持）',
    }
    label = LABELS.get(key, f'図：{key}')

    # グレー背景のプレースホルダー段落
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)

    # 枠線のような見た目にするためにXMLで網掛け設定
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')  # グレー背景
    pPr.append(shd)

    run = p.add_run(f'▼ {label} ▼\n（ここにExcelの図を貼り付けてください）')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

def add_page_number(paragraph):
    """フッター段落にページ番号フィールドを追加"""
    run = paragraph.add_run()
    for tag, attr, val in [
        ('w:fldChar', 'w:fldCharType', 'begin'),
        ('w:instrText', 'xml:space', 'preserve'),
        ('w:fldChar', 'w:fldCharType', 'end'),
    ]:
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.text = 'PAGE'
        el.set(qn(attr), val)
        run._r.append(el)

def build_document(items):
    """Wordドキュメントを構築して返す"""
    doc = Document()

    # ページ設定（A4）
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    set_default_font(doc)

    # フッター（ページ番号）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)

    # タイトルページ
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(16)
    title_run = title_p.add_run('観音台一丁目第2自治会 年報 2025年度')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x17, 0x37, 0x5E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(60)
    sub_run = sub_p.add_run('第２部　住民アンケート記録')
    sub_run.bold = True
    sub_run.font.size = Pt(20)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(4)
    date_p.add_run('2026年3月').font.size = Pt(11)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.add_run('観音台一丁目第２自治会　事務局・高橋正剛').font.size = Pt(10)

    doc.add_page_break()

    # 本文
    for item in items:
        t = item['type']
        if t == 'h2':
            add_heading(doc, item['content'], 2, 14)
        elif t == 'h3':
            add_heading(doc, item['content'], 3, 13)
        elif t == 'h4':
            add_heading(doc, item['content'], 4, 12)
        elif t == 'para':
            # 複数行の場合は行ごとに段落を分割
            lines = item['content'].split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('・'):
                    add_para(doc, line, indent=True, space_after=3)
                elif line:
                    add_para(doc, line, space_after=6)
        elif t == 'chart':
            add_chart_placeholder(doc, item['key'])

    return doc

# ============================================================
# メイン処理
# ============================================================
def main():
    print("=== 年報 第2部（アンケート編）生成スクリプト ===", flush=True)

    print(f"Markdownドラフトを読み込み中: {DRAFT_FILE}", flush=True)
    items = parse_draft(DRAFT_FILE)
    print(f"  → {len(items)} 要素を解析", flush=True)

    print("Word文書を生成中...", flush=True)
    doc = build_document(items)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"✅ 出力完了: {OUTPUT_FILE}", flush=True)

if __name__ == '__main__':
    main()
