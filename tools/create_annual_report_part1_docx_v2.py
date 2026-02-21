r"""
年報 第1部 Word文書生成スクリプト v2

draft_v2.mdから本文を読み込んでWordファイルとして出力します。

Requirements:
- python-docx (already installed)

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_part1_docx_v2.py

⚠️ CONTENT SOURCE:
This script reads the draft content from:
docs/community/2026__/annual_report_part1_draft_v2.md

To update the content, edit the markdown file and re-run this script.

⚠️ CRITICAL: ページ番号設定の注意事項
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
python-docxでセクションのページ番号を1から開始する設定は、以下の点に注意が必要：

【問題の本質】
- doc.add_section()でセクションを作成すると、そのセクションのプロパティ（sectPr）は
  セクションの**終わり**に配置される（Wordの内部XML構造の仕様）
- セクション作成直後に section._sectPr を設定しても、文書構造上の適切な位置にない

【正しい実装方法】
- セクションのページ番号開始設定は、**次のセクションを作る直前**に行う必要がある
- つまり、本文セクションのページ番号を1から始めるには：
  1. 目次セクションの終わりで本文セクションを作成
  2. 本文の内容を追加
  3. 第2部セクションを作成する**直前**に、current_section._sectPr にページ番号設定を追加
  4. 第2部セクションを作成

【誤った実装例（動作しない）】
    section_break2 = doc.add_section()  # 本文セクション作成
    section_break2._sectPr に pgNumType を設定  # ✗ 効果なし
    # ... 本文追加

【正しい実装例（動作する）】
    section_break2 = doc.add_section()  # 本文セクション作成
    # ... 本文追加
    current_section = doc.sections[-1]  # 本文セクション
    current_section._sectPr に pgNumType を設定  # ✓ 次のセクション作成前に設定
    section_break_part2 = doc.add_section()  # 第2部セクション作成

参考：2026-02-11 に発生した問題と解決策
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import sys
import re
from pathlib import Path

def extract_draft_content(md_file_path):
    """
    Markdownファイルから「## ドラフト全文」セクションを抽出
    
    ⚠️ CRITICAL: 段落分離の注意事項
    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    【問題の事例】2026-02-11に発生
    - Markdownの引用ブロック内で `>` だけの空行を段落区切りとして処理していなかった
    - 結果：複数の段落が1つの大きな段落に結合された
    - 影響：特定の段落のみに適用すべき黄色背景やインデントが、結合された段落全体に適用された
    
    【正しい実装】
    - `>` だけの空行（内容が空文字列）は段落の区切りとして処理する
    - `> 内容` の行は同じ段落として結合する
    - 通常の空行も段落の区切りとして処理する
    
    【Markdownの例】
    > 段落1の内容
    > 段落1の続き
    >
    > 段落2の内容  ← これは別の段落として処理される
    
    この処理を誤ると、太字検出による特別なフォーマット（黄色背景、インデント）が
    意図しない段落にまで適用されてしまう。
    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    """
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 「## ドラフト本文」セクションを探す（v2形式）
    # 「## 執筆メモ」の前まで全てを抽出
    pattern = r'## ドラフト本文.*?\n\n(.*?)(?=\n## 執筆メモ|\n## v1からの|$)'
    match = re.search(pattern, content, re.DOTALL)
    
    if not match:
        raise ValueError("「## ドラフト本文」セクションが見つかりません")
    
    draft_text = match.group(1)
    
    # パースして構造化
    lines = draft_text.split('\n')
    sections = []
    current_para = []
    
    for line in lines:
        stripped = line.strip()
        
        # セクション区切り（---）は無視
        if stripped == '---':
            if current_para:
                sections.append({'type': '段落', 'content': '\n'.join(current_para)})
                current_para = []
            continue
        
        # 見出し（###）
        if stripped.startswith('###'):
            if current_para:
                sections.append({'type': '段落', 'content': '\n'.join(current_para)})
                current_para = []
            sections.append({'type': '見出し', 'content': stripped.replace('###', '').strip()})
        
        # 引用ブロック（>）または通常の行
        elif stripped.startswith('>'):
            text = stripped[1:].strip()
            if text:
                current_para.append(text)
            elif current_para:
                # `>`だけの空行は段落の区切り
                sections.append({'type': '段落', 'content': '\n'.join(current_para)})
                current_para = []
        elif stripped:
            current_para.append(stripped)
        
        # 空行（段落区切り）
        elif current_para:
            sections.append({'type': '段落', 'content': '\n'.join(current_para)})
            current_para = []
    
    # 最後の段落
    if current_para:
        sections.append({'type': '段落', 'content': '\n'.join(current_para)})
    
    return sections

def parse_bold_text(text):
    """
    **bold**を認識して、(text, is_bold)のリストを返す
    """
    parts = []
    pattern = r'\*\*(.*?)\*\*'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # 太字の前の通常テキスト
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        # 太字テキスト
        parts.append((match.group(1), True))
        last_end = match.end()
    
    # 残りの通常テキスト
    if last_end < len(text):
        parts.append((text[last_end:], False))
    
    return parts if parts else [(text, False)]


def extract_second_half(md_path):
    """
    annual_report_part1_draft_v2_second.md から ドラフト本文 を抽出し
    段落・見出し・画像・テーブル の構造リストに変換する
    """
    text_content = md_path.read_text(encoding='utf-8')
    m = re.search(r'## ドラフト本文.*?\n\n(.*?)(?=\n## 執筆メモ|\Z)', text_content, re.DOTALL)
    if not m:
        return []

    lines = m.group(1).split('\n')
    items = []
    buf = []

    def flush():
        if buf:
            items.append({'type': 'para', 'content': '\n'.join(buf)})
            buf.clear()

    for line in lines:
        s = line.strip()
        if s == '---':
            flush()
        elif s.startswith('### '):
            flush()
            items.append({'type': 'heading', 'content': s[4:]})
        elif s.startswith('|') and len(s) > 1:
            flush()
            items.append({'type': 'table_row', 'content': s})
        elif s.startswith('!['):
            flush()
            img_m = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', s)
            if img_m:
                items.append({'type': 'image', 'alt': img_m.group(1), 'src': img_m.group(2)})
        elif s.startswith('**[図：'):
            flush()
            items.append({'type': 'fig_placeholder', 'content': s})
        elif s:
            buf.append(s)
        else:
            flush()

    flush()

    # 隣接する table_row をまとめる
    merged = []
    i = 0
    while i < len(items):
        if items[i]['type'] == 'table_row':
            rows = []
            while i < len(items) and items[i]['type'] == 'table_row':
                rows.append(items[i]['content'])
                i += 1
            merged.append({'type': 'table', 'rows': rows})
        else:
            merged.append(items[i])
            i += 1
    return merged


def render_md_table(doc, rows, parse_bold_fn):
    """Markdown テーブル行リストを Word テーブルに変換"""
    data = []
    for row in rows:
        if re.match(r'\|[\s\-:|]+\|', row):
            continue  # 区切り行スキップ
        cells = [c.strip() for c in row.strip('|').split('|')]
        data.append(cells)
    if not data:
        return
    ncols = max(len(r) for r in data)
    tbl = doc.add_table(rows=len(data), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            if ci >= ncols:
                break
            cell = tbl.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            for t, bold in parse_bold_fn(cell_text):
                r = p.add_run(t)
                r.bold = bold or (ri == 0)
                r.font.size = Pt(10)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    print("python-docx imported successfully", flush=True)
    
    # Markdownファイルを読み込み
    md_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'annual_report_part1_draft_v2.md'
    print(f"Reading draft from: {md_path}", flush=True)
    
    draft_sections = extract_draft_content(md_path)
    print(f"Extracted {len(draft_sections)} sections", flush=True)
    
    # 文書作成
    doc = Document()
    
    # ページ設定
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # ページ番号フィールドを追加する関数
    def add_page_number(paragraph):
        """フッター段落にページ番号フィールドを追加"""
        run = paragraph.add_run()
        
        # ページ番号フィールドのXMLを作成
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        return run
    
    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'メイリオ'
    font.size = Pt(11)
    
    # =========================
    # 表紙（タイトル+背景画像）
    # =========================
    cover_image_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'images' / 'report_cover.jpg'
    
    if cover_image_path.exists():
        # 表紙セクションの余白を0に設定（フチなし印刷用）
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        
        # 上部の空白
        spacer1 = doc.add_paragraph()
        spacer1.paragraph_format.space_after = Pt(50)
        
        # 自治会名
        assoc_name = doc.add_paragraph()
        assoc_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assoc_name.paragraph_format.space_after = Pt(20)
        run1 = assoc_name.add_run('観音台第二自治会')
        run1.font.size = Pt(40)
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0, 51, 102)  # 濃紺
        
        # 年報タイトル
        report_title = doc.add_paragraph()
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title.paragraph_format.space_after = Pt(15)
        run2 = report_title.add_run('年報 2025（案）')
        run2.font.size = Pt(56)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0, 51, 102)  # 濃紺
        
        # キャッチフレーズ：Looking Forward / with / Shared Diversity
        catchphrase1 = doc.add_paragraph()
        catchphrase1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        catchphrase1.paragraph_format.space_after = Pt(4)
        run3a = catchphrase1.add_run('Looking Forward')
        run3a.font.size = Pt(26)
        run3a.font.italic = True
        run3a.font.color.rgb = RGBColor(70, 130, 180)  # スチールブルー

        catchphrase2 = doc.add_paragraph()
        catchphrase2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        catchphrase2.paragraph_format.space_after = Pt(4)
        run3b = catchphrase2.add_run('with')
        run3b.font.size = Pt(15)
        run3b.font.italic = True
        run3b.font.color.rgb = RGBColor(100, 149, 190)  # やや薄め

        catchphrase3 = doc.add_paragraph()
        catchphrase3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        catchphrase3.paragraph_format.space_after = Pt(40)
        run3c = catchphrase3.add_run('Shared Diversity')
        run3c.font.size = Pt(26)
        run3c.font.italic = True
        run3c.font.color.rgb = RGBColor(70, 130, 180)  # スチールブルー
        
        # 表紙：画像を配置
        cover_para = doc.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_para.paragraph_format.space_before = Pt(0)
        cover_para.paragraph_format.space_after = Pt(0)
        
        run = cover_para.add_run()
        # 画像サイズをA4ページ全体に拡大（余白なし）
        picture = run.add_picture(str(cover_image_path), width=Inches(8.27), height=Inches(11.69))
        
        # 表紙の注記
        cover_note = doc.add_paragraph()
        cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_note.paragraph_format.space_before = Pt(30)
        cover_note.paragraph_format.space_after = Pt(20)
        note_run = cover_note.add_run('これは、案として、およその、年報のイメージをお伝えするもので、\n実際のものは、さらに内容についての、検討を加えたものになります。')
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(102, 102, 102)  # グレー
        note_run.font.italic = True
        
        # セクション区切りを追加（表紙セクション終了）
        section_break = doc.add_section()
        section_break.start_type = 2  # 次のページから新しいセクション
        
        # 目次セクションの余白を1インチに設定（表紙の余白0設定をリセット）
        section_break.top_margin = Inches(1.0)
        section_break.bottom_margin = Inches(1.0)
        section_break.left_margin = Inches(1.0)
        section_break.right_margin = Inches(1.0)
        
        # 目次セクションのフッターもページ番号なし
        section_break.footer.is_linked_to_previous = False
        section_break.footer.paragraphs[0].clear()
    
    # =========================
    # 目次ページ
    # =========================
    
    # タイトル
    toc_title = doc.add_heading('目次', level=1)
    toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(30)
    
    # 第1部
    part1_heading = doc.add_paragraph()
    part1_heading.paragraph_format.space_before = Pt(10)
    part1_heading.paragraph_format.space_after = Pt(10)
    run_p1 = part1_heading.add_run('第1部：考え方と風土の共有')
    run_p1.bold = True
    run_p1.font.size = Pt(14)
    run_p1.font.color.rgb = RGBColor(0, 51, 102)
    
    # 第1部の項目
    items_part1 = [
        '1. 導入',
        '2. 構造の発見',
        '3. 矛盾の具体化',
        '4. 健康な構造のビジョン',
        '5. 展望と問いの共有'
    ]
    
    for item in items_part1:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(11)
    
    # 空白
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(20)
    
    # 第2部
    part2_heading = doc.add_paragraph()
    part2_heading.paragraph_format.space_before = Pt(10)
    part2_heading.paragraph_format.space_after = Pt(10)
    run_p2 = part2_heading.add_run('第2部：総会資料')
    run_p2.bold = True
    run_p2.font.size = Pt(14)
    run_p2.font.color.rgb = RGBColor(0, 51, 102)
    
    # 第2部の項目
    items_part2 = [
        '第1号議案：2025年度事業報告',
        '第2号議案：2025年度会計報告',
        '第3号議案：2026年度役員改選案',
        '第4号議案：2026年度事業計画案',
        '第5号議案：2026年度予算案',
        '第6号議案：会則の一部改正'
    ]
    
    for item in items_part2:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(11)
    
    # 参考資料
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_after = Pt(10)
    
    ref_heading = doc.add_paragraph()
    ref_heading.paragraph_format.space_before = Pt(10)
    ref_heading.paragraph_format.space_after = Pt(10)
    run_ref = ref_heading.add_run('参考資料')
    run_ref.bold = True
    run_ref.font.size = Pt(14)
    run_ref.font.color.rgb = RGBColor(0, 51, 102)
    
    items_ref = [
        '・会則全体',
        '・会員名簿',
        '・会員住居地図'
    ]
    
    for item in items_ref:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(11)
    
    # セクション区切り（目次セクション終了）
    section_break2 = doc.add_section()
    section_break2.start_type = 2  # 次のページから新しいセクション
    
    # 本文セクションの余白を1インチに設定
    section_break2.top_margin = Inches(1.0)
    section_break2.bottom_margin = Inches(1.0)
    section_break2.left_margin = Inches(1.0)
    section_break2.right_margin = Inches(1.0)
    
    # 本文セクションのフッターにページ番号を設定
    section_break2.footer.is_linked_to_previous = False
    footer_para = section_break2.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)
    
    # =========================
    # 本文開始（draft_v2.mdから読み込み）
    # =========================
    
    subtitle = doc.add_heading('第1部：考え方と風土の共有', level=2)
    subtitle_format = subtitle.paragraph_format
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format.space_after = Pt(20)
    subtitle_format.space_before = Pt(10)
    subtitle_format.left_indent = Inches(0)
    subtitle_format.first_line_indent = Inches(0)
    
    # Markdownから抽出したセクションをWordに変換
    hierarchy_diagram_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'images' / 'hierarchy_diagram.png'
    after_structure_heading = False
    inserted_diagram = False
    
    for sect in draft_sections:
        if sect['type'] == '見出し':
            # 見出し
            heading = doc.add_heading(sect['content'], level=3)
            heading.paragraph_format.space_after = Pt(12)
            heading.paragraph_format.left_indent = Inches(0)
            heading.paragraph_format.first_line_indent = Inches(0)
            
            # 「構造の発見」の見出しを検出
            if '構造の発見' in sect['content']:
                after_structure_heading = True
        
        elif sect['type'] == '段落':
            # 段落
            para = doc.add_paragraph(style='Normal')
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.2
            para.paragraph_format.left_indent = Inches(0)
            para.paragraph_format.first_line_indent = Inches(0)
            
            # 太字を認識してテキストを追加
            text_parts = parse_bold_text(sect['content'])
            for text, is_bold in text_parts:
                run = para.add_run(text)
                if is_bold:
                    run.bold = True
                    # 強調ボックススタイル（黄色背景）
                    if '意見が分かれた理由' in text or '何を求めているか' in text:
                        para.paragraph_format.left_indent = Inches(0.3)
                        para.paragraph_format.right_indent = Inches(0.3)
                        # 背景色
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'FFF9E6')
                        para._element.get_or_add_pPr().append(shading)
            
            # 「構造の発見」セクションで図の参照を検出したら図を挿入
            if after_structure_heading and not inserted_diagram:
                # v2では「**[図：自治会への期待の変化]**」という行がある
                if '[図：' in sect['content'] or '図：' in sect['content']:
                    if hierarchy_diagram_path.exists():
                        spacer_before_img = doc.add_paragraph()
                        spacer_before_img.paragraph_format.space_after = Pt(5)
                        
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.paragraph_format.space_before = Pt(10)
                        img_para.paragraph_format.space_after = Pt(10)
                        img_run = img_para.add_run()
                        img_run.add_picture(str(hierarchy_diagram_path), width=Inches(6.5))
                        
                        # 図のキャプション
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.paragraph_format.space_after = Pt(15)
                        caption_run = caption.add_run('図：自治会への期待：時代による変化')
                        caption_run.font.size = Pt(9)
                        caption_run.italic = True
                        caption_run.font.color.rgb = RGBColor(102, 102, 102)
                        
                        # フラグを更新
                        inserted_diagram = True
                    continue  # 図の参照行自体は表示しない
    
    # =========================
    # 第1部（後半）：財政の30年見通し
    # =========================
    md_second_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'annual_report_part1_draft_v2_second.md'
    garbage_graph_path = Path(__file__).parent / 'forecast_garbage_cumulative.png'

    if md_second_path.exists():
        doc.add_page_break()

        second_subtitle = doc.add_heading('第1部（続き）：数字で見る30年の見通し', level=2)
        second_subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        second_subtitle.paragraph_format.space_after = Pt(20)
        second_subtitle.paragraph_format.left_indent = Inches(0)
        second_subtitle.paragraph_format.first_line_indent = Inches(0)

        second_sections = extract_second_half(md_second_path)
        inserted_graph = False

        for item in second_sections:
            if item['type'] == 'heading':
                h = doc.add_heading(item['content'], level=3)
                h.paragraph_format.space_after = Pt(10)
                h.paragraph_format.left_indent = Inches(0)
                h.paragraph_format.first_line_indent = Inches(0)
            elif item['type'] == 'para':
                para = doc.add_paragraph(style='Normal')
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.line_spacing = 1.2
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.first_line_indent = Inches(0)
                for text_part, is_bold in parse_bold_text(item['content']):
                    r = para.add_run(text_part)
                    if is_bold:
                        r.bold = True
            elif item['type'] in ('image', 'fig_placeholder'):
                src = item.get('src', item.get('content', ''))
                if 'forecast_garbage' in src or 'ゴミ' in src:
                    if not inserted_graph and garbage_graph_path.exists():
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.paragraph_format.space_before = Pt(10)
                        img_para.paragraph_format.space_after = Pt(10)
                        img_para.add_run().add_picture(str(garbage_graph_path), width=Inches(6.2))
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.paragraph_format.space_after = Pt(15)
                        cap_run = cap.add_run(
                            '図：ゴミ集積所使用料 累計推移（上：総資産推移 ／ 下：非会員数・使用料累計）'
                        )
                        cap_run.font.size = Pt(9)
                        cap_run.italic = True
                        cap_run.font.color.rgb = RGBColor(102, 102, 102)
                        inserted_graph = True
            elif item['type'] == 'table':
                render_md_table(doc, item['rows'], parse_bold_text)
    else:
        print(f'  ⚠ 後半ファイルが見つかりません: {md_second_path}', flush=True)

    # =========================
    # フッター
    # =========================
    footer = doc.add_paragraph()
    footer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(20)
    footer_run = footer.add_run('2026年2月\n観音台自治会 事務局')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Attribution
    attribution = doc.add_paragraph()
    attribution.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    attribution.paragraph_format.space_before = Pt(5)
    attr_run = attribution.add_run('Draft created with assistance from GitHub Copilot')
    attr_run.font.size = Pt(8)
    attr_run.font.italic = True
    attr_run.font.color.rgb = RGBColor(153, 153, 153)
    
    # =========================
    # 第2部の構成概要ページ
    # =========================
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # ⚠️ CRITICAL: ページ番号を1から開始する設定
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 【重要】この設定は、次のセクション（第2部）を作成する**直前**に行う必要がある。
    # 
    # python-docxでは、セクションのプロパティ（sectPr）はセクションの**終わり**に配置される。
    # そのため、セクション作成直後に設定しても効果がない。
    # 
    # 正しい順序：
    # 1. 本文セクション作成（section_break2 = doc.add_section()）←すでに実行済み
    # 2. 本文コンテンツ追加（draft_sectionsのループ）←すでに実行済み
    # 3. 本文セクションのsectPrにページ番号設定を追加 ←ここで実行
    # 4. 次のセクション（第2部）を作成 ←この後すぐ実行
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    
    # 第2部のセクションを追加する前に、本文セクションのページ番号を1から開始するように設定
    current_section = doc.sections[-1]  # 現在のセクション = 本文セクション
    sectPr = current_section._sectPr
    # 既存のpgNumType要素を削除
    for pgNumType_elem in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(pgNumType_elem)
    # 新しいpgNumType要素を追加
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.insert(0, pgNumType)
    
    section_break_part2 = doc.add_section()
    section_break_part2.start_type = 2
    section_break_part2.footer.is_linked_to_previous = True  # 第1部からのページ番号を継続
    
    # タイトル
    part2_visual_title = doc.add_heading('第2部：総会資料（構成イメージ）', level=2)
    part2_visual_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part2_visual_title.paragraph_format.space_after = Pt(10)
    
    # 説明文
    explanation = doc.add_paragraph()
    explanation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    explanation.paragraph_format.space_after = Pt(20)
    exp_run = explanation.add_run('※ 以下は第2部の構成イメージです。実際の内容は従来通りの総会資料となります。')
    exp_run.font.size = Pt(9)
    exp_run.font.color.rgb = RGBColor(102, 102, 102)
    exp_run.italic = True
    
    # 議案データ
    agendas = [
        ('第1号議案：2025年度事業報告', '会員の交流促進、住宅環境整備、市への要望など'),
        ('第2号議案：2025年度会計報告', '収入・支出の詳細、監査報告'),
        ('第3号議案：2026年度役員改選案', '新役員体制の提案'),
        ('第4号議案：2026年度事業計画案', '今年度の活動計画'),
        ('第5号議案：2026年度予算案', '収入・支出の予算'),
        ('第6号議案：会則の一部改正', '会則改正の提案と理由'),
        ('参考資料', '会則全体、会員名簿、会員住居地図')
    ]
    
    # テーブルで各項目を枠で囲む
    for agenda_title, agenda_desc in agendas:
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.rows[0].cells[0]
        
        # セルの背景色
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F0F8FF')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # セルの余白
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '150')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        # セル内コンテンツ
        cell_para = cell.paragraphs[0]
        
        run1 = cell_para.add_run(agenda_title)
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = RGBColor(0, 51, 102)
        
        cell_para.add_run('\n')
        
        run2 = cell_para.add_run(agenda_desc)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(80, 80, 80)
        
        table.alignment = 1
        table.width = Inches(6.5)
        
        # 枠線の色
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:color'), '2C5F8D')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # テーブル間のスペース
        spacer_para = doc.add_paragraph()
        spacer_para.paragraph_format.space_after = Pt(15)
    
    # =========================
    # 裏表紙
    # =========================
    if cover_image_path.exists():
        # 裏表紙セクションを作る前に、第2部のページ番号継続を確保
        part2_section = doc.sections[-1]
        part2_sectPr = part2_section._sectPr
        # pgNumType要素があれば削除（ページ番号リセットを防ぐ）
        for pgNumType_elem in part2_sectPr.findall(qn('w:pgNumType')):
            part2_sectPr.remove(pgNumType_elem)
        
        section_break3 = doc.add_section()
        section_break3.start_type = 2
        
        # 裏表紙セクションの余白を0に設定（フチなし印刷用）
        section_break3.top_margin = Inches(0)
        section_break3.bottom_margin = Inches(0)
        section_break3.left_margin = Inches(0)
        section_break3.right_margin = Inches(0)
        
        # 裏表紙セクションのフッターはページ番号なし
        section_break3.footer.is_linked_to_previous = False
        section_break3.footer.paragraphs[0].clear()
        
        back_cover_paragraph = doc.add_paragraph()
        back_cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        back_cover_run = back_cover_paragraph.add_run()
        # 画像サイズをA4ページ全体に拡大（余白なし）
        back_cover_run.add_picture(str(cover_image_path), width=Inches(8.27), height=Inches(11.69))
        back_cover_paragraph.paragraph_format.space_after = Pt(0)
        back_cover_paragraph.paragraph_format.space_before = Pt(0)
    
    # 保存
    output_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'annual_report_part1_v2.docx'
    doc.save(output_path)
    
    print(f"✓ SUCCESS! Word文書を作成しました", flush=True)
    print(f"✓ Output: {output_path.relative_to(Path.cwd())}", flush=True)
    print(f"✓ File size: {output_path.stat().st_size:,} bytes", flush=True)
    
except ImportError as e:
    print(f"ERROR: Missing required library - {e}", file=sys.stderr, flush=True)
    print("\nInstall missing libraries:", flush=True)
    print('  & "C:\\Program Files\\Python313\\python.exe" -m pip install python-docx', flush=True)
    sys.exit(1)
    
except Exception as e:
    print(f"ERROR: {e}", file=sys.stderr, flush=True)
    import traceback
    traceback.print_exc()
    sys.exit(1)
