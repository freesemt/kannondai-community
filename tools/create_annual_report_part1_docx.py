"""
年報 第1部 Word文書生成スクリプト

draft_v1の本文をWordファイルとして出力します。

Requirements:
- python-docx (already installed)

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_part1_docx.py

⚠️ IMPORTANT - Cover Page Design Limitation:

python-docx cannot directly overlay text on images as background.
Attempted approaches that failed:
- Negative spacing (value range error)
- Header-based background (images appeared on all pages)
- Direct XML/DrawingML manipulation (file corruption)

Current approach: Simple layout (title + image) for reliable generation.
Manual adjustment in Word required for text overlay effect.

Manual Steps:
1. Right-click image → Text Wrapping → Behind Text
2. Insert → Text Box → Add title
3. Position text box over image
4. Adjust font, color, positioning

This approach is intentional and documented in tools/README.md.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    print("python-docx imported successfully", flush=True)
    
    # 文書作成
    doc = Document()
    
    # ページ設定
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # ページ番号フィールドを追加する関数
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def add_page_number(paragraph):
        """フッター段落にページ番号フィールドを追加"""
        run = paragraph.add_run()
        
        # ページ番号フィールドのXMLを作成
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        return run
    
    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'メイリオ'
    font.size = Pt(11)
    
    # =========================
    # 表紙（画像とタイトルを配置、手動で調整可能に）
    # =========================
    cover_image_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'images' / 'report_cover.jpg'
    
    if cover_image_path.exists():
        # 上部の空白
        spacer1 = doc.add_paragraph()
        spacer1.paragraph_format.space_after = Pt(50)
        
        # 自治会名
        assoc_name = doc.add_paragraph()
        assoc_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assoc_name.paragraph_format.space_after = Pt(20)
        run1 = assoc_name.add_run('観音台第二自治会')
        run1.font.size = Pt(40)
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0, 51, 102)  # 濃紺
        
        # 年報タイトル
        report_title = doc.add_paragraph()
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title.paragraph_format.space_after = Pt(40)
        run2 = report_title.add_run('年報 2025（案）')
        run2.font.size = Pt(56)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0, 51, 102)  # 濃紺
        
        # 表紙：画像を配置
        cover_para = doc.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_para.paragraph_format.space_before = Pt(0)
        cover_para.paragraph_format.space_after = Pt(0)
        
        run = cover_para.add_run()
        # 画像サイズを調整
        picture = run.add_picture(str(cover_image_path), width=Inches(6.0))
        
        # セクション区切りを追加（表紙セクション終了）
        # 改ページではなくセクション区切り
        section_break = doc.add_section()
        section_break.start_type = 2  # 次のページから新しいセクション
        
        # 目次セクションのフッターもページ番号なし（前のセクションとリンクしない）
        section_break.footer.is_linked_to_previous = False
        section_break.footer.paragraphs[0].clear()
    
    # =========================
    # 目次ページ
    # =========================
    
    # タイトル
    toc_title = doc.add_heading('目次', level=1)
    toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(30)
    
    # 第1部
    part1_heading = doc.add_paragraph()
    part1_heading.paragraph_format.space_before = Pt(10)
    part1_heading.paragraph_format.space_after = Pt(10)
    run_p1 = part1_heading.add_run('第1部：考え方と風土の共有')
    run_p1.bold = True
    run_p1.font.size = Pt(14)
    run_p1.font.color.rgb = RGBColor(0, 51, 102)
    
    # 第1部の項目
    items_part1 = [
        '・役員引受の悩みと構造的問題',
        '・自治会に期待する階層（Level 1-3）',
        '・なぜ気づきにくかったのか',
        '・今年度の取り組み',
        '・来年度へ'
    ]
    
    for item in items_part1:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(11)
    
    # 空白
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(20)
    
    # 第2部
    part2_heading = doc.add_paragraph()
    part2_heading.paragraph_format.space_before = Pt(10)
    part2_heading.paragraph_format.space_after = Pt(10)
    run_p2 = part2_heading.add_run('第2部：総会資料')
    run_p2.bold = True
    run_p2.font.size = Pt(14)
    run_p2.font.color.rgb = RGBColor(0, 51, 102)
    
    # 第2部の項目（議案）
    items_part2 = [
        '第1号議案：2025年度事業報告',
        '第2号議案：2025年度会計報告',
        '第3号議案：2026年度役員改選案',
        '第4号議案：2026年度事業計画案',
        '第5号議案：2026年度予算案',
        '第6号議案：会則の一部改正'
    ]
    
    for item in items_part2:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(11)
    
    # 参考資料
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_after = Pt(10)
    
    ref_heading = doc.add_paragraph()
    ref_heading.paragraph_format.space_before = Pt(10)
    ref_heading.paragraph_format.space_after = Pt(10)
    run_ref = ref_heading.add_run('参考資料')
    run_ref.bold = True
    run_ref.font.size = Pt(14)
    run_ref.font.color.rgb = RGBColor(0, 51, 102)
    
    items_ref = [
        '・会則全体',
        '・会員名簿',
        '・会員住居地図'
    ]
    
    for item in items_ref:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(11)
    
    # セクション区切りを追加（目次セクション終了）
    section_break2 = doc.add_section()
    section_break2.start_type = 2  # 次のページから新しいセクション
    
    # 本文セクションのフッターにページ番号を設定
    # 前のセクションのフッターとリンクしない
    section_break2.footer.is_linked_to_previous = False
    footer_para = section_break2.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)
    
    # =========================
    # 本文開始（タイトルなし、直接第1部から）
    # =========================
    
    subtitle = doc.add_heading('第1部：考え方と風土の共有', level=2)
    subtitle_format = subtitle.paragraph_format
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format.space_after = Pt(20)
    subtitle_format.space_before = Pt(10)
    
    # =========================
    # 本文
    # =========================
    
    # 第1段落：春の木の比喩
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(12)
    p1.paragraph_format.line_spacing = 1.5
    p1.add_run('春が来ると、庭の木々は一斉に芽吹きます。でも、よく見ると、それぞれの木が違うタイミング、違う速さで育っています。')
    
    # 第2段落：自治会の多様性
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.5
    p2.add_run('自治会も同じかもしれません。112世帯、それぞれの家庭に、それぞれの事情があります。今年度、私たちは、この「それぞれの事情」を、少しだけ丁寧に見つめようとしました。')
    
    # 第3段落：アンケート結果
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(12)
    p3.paragraph_format.line_spacing = 1.5
    p3.add_run('昨年は夏と冬に２回、アンケートを実施しました。冬には、89世帯から回答をいただきました。会費改定について、意見は分かれました。現状維持、簡便性重視、多様化許容。三つの選択肢に、それぞれ支持がありました。')
    
    # 第4段落：重要なことが見えてくる
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(12)
    p4.paragraph_format.line_spacing = 1.5
    p4.add_run('それを含めて、この一年を振り返ると、重要なことが見えてきます。')
    
    # 第5段落：根本的な違い（強調ボックス）
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(12)
    p5.paragraph_format.line_spacing = 1.5
    p5.paragraph_format.left_indent = Inches(0.3)
    p5.paragraph_format.right_indent = Inches(0.3)
    p5_run = p5.add_run('意見が分かれた理由は、いろいろ考えられます。問い方がよくなかった点もあります。しかし、それぞれの家庭が、自治会に「何を求めているか」が、個々の事情によって違っている点が大きいのではないでしょうか。')
    p5_run.bold = True
    # 背景色（薄い黄色）
    shading_p5 = OxmlElement('w:shd')
    shading_p5.set(qn('w:fill'), 'FFF9E6')
    p5._element.get_or_add_pPr().append(shading_p5)
    
    # 第6段落：ニーズの多様性
    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(12)
    p6.paragraph_format.line_spacing = 1.5
    p6.add_run('ある家庭は、ゴミ出しができればよい。ある家庭は、防災も大切にしたい。ある家庭は、もっと別のことを求めているかもしれません。')
    
    # 第7段落：構造の出現
    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(20)
    p7.paragraph_format.line_spacing = 1.5
    p7.add_run('この「何を求めているか」の違いと、現在の仕組み。この二つを並べて見ると、ある構造が浮かび上がってきます。')
    
    # =========================
    # セクション：その構造とは
    # =========================
    section1 = doc.add_heading('その構造とは', level=3)
    section1.paragraph_format.space_after = Pt(12)
    
    # 導入：三つの階層
    p8 = doc.add_paragraph()
    p8.paragraph_format.space_after = Pt(12)
    p8.paragraph_format.line_spacing = 1.5
    p8.add_run('話を簡単にするため、自治会に期待することは、大きく三つの階層に分けられると仮定してみます。')
    
    # 第一の階層
    p9 = doc.add_paragraph()
    p9.paragraph_format.space_after = Pt(8)
    p9.paragraph_format.line_spacing = 1.5
    p9_run1 = p9.add_run('第一の階層')
    p9_run1.bold = True
    p9.add_run('は、普通に生きていくための最低限のこと。ゴミを捨てられること。これは誰にとっても必要です。選択の余地はありません。')
    
    # 第二の階層
    p10 = doc.add_paragraph()
    p10.paragraph_format.space_after = Pt(8)
    p10.paragraph_format.line_spacing = 1.5
    p10_run1 = p10.add_run('第二の階層')
    p10_run1.bold = True
    p10.add_run('は、安全や安心。防災、防犯、見守り。これも大切ですが、どこまで期待するかは、少し人によって違います。')
    
    # 第三の階層
    p11 = doc.add_paragraph()
    p11.paragraph_format.space_after = Pt(12)
    p11.paragraph_format.line_spacing = 1.5
    p11_run1 = p11.add_run('第三の階層')
    p11_run1.bold = True
    p11.add_run('は、もっと充実した生活。交流イベント、コミュニティの活動。これは人によって、価値の感じ方が大きく異なります。')
    
    # 現在の仕組み
    p12 = doc.add_paragraph()
    p12.paragraph_format.space_after = Pt(12)
    p12.paragraph_format.line_spacing = 1.5
    p12.add_run('現在の自治会の仕組みはどうなっているでしょうか。')
    
    p13 = doc.add_paragraph()
    p13.paragraph_format.space_after = Pt(12)
    p13.paragraph_format.line_spacing = 1.5
    p13.add_run('会員であれば、年額3,600円で、すべての階層のサービスを利用できます。そして、役員として活動に参加する義務があります。')
    
    # Aタイプの考察
    p14 = doc.add_paragraph()
    p14.paragraph_format.space_after = Pt(12)
    p14.paragraph_format.line_spacing = 1.5
    p14.add_run('しかし、ここで第一の階層（ゴミ）だけを必要とする家庭（以下、Ａタイプといいます）を考えてみましょう。誤解を避けるための補足ですが、この状況は時代によって意味が異なります。かつての貧しかった日本と、比較的豊かになった今との違いです。昔なら、第２、第３の階層もなくて、最低でも第１階層だけが欲しいという意味になりますが、今は違います。第２、第３の階層は自分で充足できるので、第１階層だけが欲しいのです。')
    
    p15 = doc.add_paragraph()
    p15.paragraph_format.space_after = Pt(12)
    p15.paragraph_format.line_spacing = 1.5
    p15.add_run('さて、Ａタイプの場合の自治会への参加はどうなるでしょうか？次の２つの選択が考えられます。')
    
    # 選択肢（インデント付き）
    p16 = doc.add_paragraph()
    p16.paragraph_format.space_after = Pt(3)
    p16.paragraph_format.line_spacing = 1.5
    p16.paragraph_format.left_indent = Inches(0.5)
    p16_run1 = p16.add_run('選択肢１：')
    p16_run1.bold = True
    p16.add_run('第２、第３階層は不要なのに、セット販売なので、矛盾を感じながらも全サービスを買う')
    
    p17 = doc.add_paragraph()
    p17.paragraph_format.space_after = Pt(12)
    p17.paragraph_format.line_spacing = 1.5
    p17.paragraph_format.left_indent = Inches(0.5)
    p17_run1 = p17.add_run('選択肢２：')
    p17_run1.bold = True
    p17.add_run('高くなる（年額12,000円）が、信条である簡素第一主義を守って、第１階層だけを買う')
    
    p18 = doc.add_paragraph()
    p18.paragraph_format.space_after = Pt(12)
    p18.paragraph_format.line_spacing = 1.5
    p18.add_run('どちらの選択もおかしくないですか？一方は、不要なものまで買うという無駄。他方は得られるサービスが少ないのになぜか高い。')
    
    p19 = doc.add_paragraph()
    p19.paragraph_format.space_after = Pt(12)
    p19.paragraph_format.line_spacing = 1.5
    p19.add_run('ここで、「義務の方の話はどうなってるの」って思いましたか？一遍に話すと混乱しやすいので、まず一旦、切り離して見てみたということです。')
    
    p20 = doc.add_paragraph()
    p20.paragraph_format.space_after = Pt(20)
    p20.paragraph_format.line_spacing = 1.5
    p20.add_run('まずここでの注意点は、このような構造として見たとき、すべて階層を期待する家庭（以下、Cタイプといいます）が、無意識のうちに、Ａタイプ家庭に対して他の２階層のサービスを強要していないか？です。')
    
    # =========================
    # セクション：なぜ、気づきにくかったのか
    # =========================
    section2 = doc.add_heading('なぜ、気づきにくかったのか', level=3)
    section2.paragraph_format.space_after = Pt(12)
    
    p21 = doc.add_paragraph()
    p21.paragraph_format.space_after = Pt(12)
    p21.paragraph_format.line_spacing = 1.5
    p21.add_run('かつて、隣近所で助け合い、地域の行事に参加することは、とても自然なことでした。自治会の活動も、イベントも、多くの家庭にとって負担ではなく、むしろ楽しみの一つだったかもしれません。')
    
    p22 = doc.add_paragraph()
    p22.paragraph_format.space_after = Pt(12)
    p22.paragraph_format.line_spacing = 1.5
    p22.add_run('そして、会費を集め、役員を分担し、イベントを企画する仕組みは、これまでと同じように続いてきました。')
    
    p23 = doc.add_paragraph()
    p23.paragraph_format.space_after = Pt(12)
    p23.paragraph_format.line_spacing = 1.5
    p23.add_run('でも、時代は少しずつ変わっています。働き方も、暮らし方も、何を大切にするかも、少しずつ多様になってきました。')
    
    p24 = doc.add_paragraph()
    p24.paragraph_format.space_after = Pt(20)
    p24.paragraph_format.line_spacing = 1.5
    p24.add_run('多数派にとって問題がないとき、少数の人の困難は、なかなか見えません。そして、社会の変化は、ゆっくりと進むため、私たち自身も気づきにくいのかもしれません。')
    
    # =========================
    # セクション：今年度の取り組み
    # =========================
    section3 = doc.add_heading('今年度の取り組み', level=3)
    section3.paragraph_format.space_after = Pt(12)
    
    p25 = doc.add_paragraph()
    p25.paragraph_format.space_after = Pt(12)
    p25.paragraph_format.line_spacing = 1.5
    p25.add_run('この一年、いくつかの変化がありました。')
    
    p26 = doc.add_paragraph()
    p26.paragraph_format.space_after = Pt(12)
    p26.paragraph_format.line_spacing = 1.5
    p26.add_run('第6班で、班長のなり手がいなくなりました。これは、単なる人手不足ではなく、根本的な問題を示しているのかもしれません。')
    
    p27 = doc.add_paragraph()
    p27.paragraph_format.space_after = Pt(20)
    p27.paragraph_format.line_spacing = 1.5
    p27.add_run('役員会では、新しい選択肢を検討しました。たとえば、「賛助会員」という仕組み。ゴミと防災だけを希望する家庭のために、年額6,000円で、役員義務なし。')
    
    p28 = doc.add_paragraph()
    p28.paragraph_format.space_after = Pt(20)
    p28.paragraph_format.line_spacing = 1.5
    p28.add_run('これは、まだ提案の段階です。でも、多様性を認める第一歩になるかもしれません。')
    
    # =========================
    # セクション：来年度へ
    # =========================
    section4 = doc.add_heading('来年度へ', level=3)
    section4.paragraph_format.space_after = Pt(12)
    
    p29 = doc.add_paragraph()
    p29.paragraph_format.space_after = Pt(12)
    p29.paragraph_format.line_spacing = 1.5
    p29.add_run('自治会は、112世帯の共同体です。それぞれの家庭に、それぞれの事情があります。')
    
    p30 = doc.add_paragraph()
    p30.paragraph_format.space_after = Pt(12)
    p30.paragraph_format.line_spacing = 1.5
    p30.add_run('これからも、一緒に考えていきたいと思います。どうすれば、誰もが自分の生き方を尊重され、同時に、必要な助け合いができるのか。')
    
    p31 = doc.add_paragraph()
    p31.paragraph_format.space_after = Pt(30)
    p31.paragraph_format.line_spacing = 1.5
    p31.add_run('答えはまだ、見つかっていません。でも、問いを共有することが、最初の一歩ではないでしょうか。')
    
    # =========================
    # フッター
    # =========================
    footer = doc.add_paragraph()
    footer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(20)
    footer_run = footer.add_run('2026年2月\n神南大自治会 事務局')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Attribution
    attribution = doc.add_paragraph()
    attribution.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    attribution.paragraph_format.space_before = Pt(5)
    attr_run = attribution.add_run('Draft created with assistance from GitHub Copilot')
    attr_run.font.size = Pt(8)
    attr_run.font.italic = True
    attr_run.font.color.rgb = RGBColor(153, 153, 153)
    
    # =========================
    # 第2部の構成概要ページ（視覚的イメージ）
    # =========================
    
    # セクション区切り（第2部の前）
    section_break_part2 = doc.add_section()
    section_break_part2.start_type = 2
    # フッターを本文セクションと同じに設定（ページ番号継続）
    section_break_part2.footer.is_linked_to_previous = True
    
    # タイトル
    part2_visual_title = doc.add_heading('第2部：総会資料（構成イメージ）', level=2)
    part2_visual_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part2_visual_title.paragraph_format.space_after = Pt(10)
    
    # 説明文
    explanation = doc.add_paragraph()
    explanation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    explanation.paragraph_format.space_after = Pt(20)
    exp_run = explanation.add_run('※ 以下は第2部の構成イメージです。実際の内容は従来通りの総会資料となります。')
    exp_run.font.size = Pt(9)
    exp_run.font.color.rgb = RGBColor(102, 102, 102)
    exp_run.italic = True
    
    # 議案データ
    agendas = [
        ('第1号議案：2025年度事業報告', '会員の交流促進、住宅環境整備、市への要望など'),
        ('第2号議案：2025年度会計報告', '収入・支出の詳細、監査報告'),
        ('第3号議案：2026年度役員改選案', '新役員体制の提案'),
        ('第4号議案：2026年度事業計画案', '今年度の活動計画'),
        ('第5号議案：2026年度予算案', '収入・支出の予算'),
        ('第6号議案：会則の一部改正', '会則改正の提案と理由'),
        ('参考資料', '会則全体、会員名簿、会員住居地図')
    ]
    
    # テーブルを使用して各項目を明確に分離
    for agenda_title, agenda_desc in agendas:
        # 1行1列のテーブル（枠として使用）
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        # セルの設定
        cell = table.rows[0].cells[0]
        
        # セルの背景色
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F0F8FF')  # 薄青
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # セルの余白設定
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '150')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        # セル内のコンテンツ
        cell_para = cell.paragraphs[0]
        
        # タイトル
        run1 = cell_para.add_run(agenda_title)
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = RGBColor(0, 51, 102)
        
        cell_para.add_run('\n')
        
        # 説明
        run2 = cell_para.add_run(agenda_desc)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(80, 80, 80)
        
        # テーブル全体の左右マージン
        table.alignment = 1  # 中央寄せ
        
        # テーブルの幅を設定
        table.width = Inches(6.5)
        
        # 枠線の色を濃紺に設定
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:color'), '2C5F8D')  # 濃紺
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # テーブル間のスペース
        spacer_para = doc.add_paragraph()
        spacer_para.paragraph_format.space_after = Pt(15)
    
    # =========================
    # 裏表紙（背景画像のみ）
    # =========================
    if cover_image_path.exists():
        # セクション区切りを追加（裏表紙セクション）
        section_break3 = doc.add_section()
        section_break3.start_type = 2  # 次のページから新しいセクション
        
        # 裏表紙セクションのフッターはページ番号なし
        section_break3.footer.is_linked_to_previous = False
        section_break3.footer.paragraphs[0].clear()
        
        # 背景画像を挿入
        back_cover_paragraph = doc.add_paragraph()
        back_cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        back_cover_run = back_cover_paragraph.add_run()
        back_cover_run.add_picture(str(cover_image_path), width=Inches(6.0))
        back_cover_paragraph.paragraph_format.space_after = Pt(0)
        back_cover_paragraph.paragraph_format.space_before = Pt(0)
    
    # 保存
    output_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'annual_report_part1_v1.docx'
    doc.save(output_path)
    
    print(f"✓ SUCCESS! Word文書を作成しました", flush=True)
    print(f"✓ Output: {output_path.relative_to(Path.cwd())}", flush=True)
    print(f"✓ File size: {output_path.stat().st_size:,} bytes", flush=True)
    
except ImportError as e:
    print(f"ERROR: Missing required library - {e}", file=sys.stderr, flush=True)
    print("\nInstall missing libraries:", flush=True)
    print('  & "C:\\Program Files\\Python313\\python.exe" -m pip install python-docx', flush=True)
    sys.exit(1)
    
except Exception as e:
    print(f"ERROR: {e}", file=sys.stderr, flush=True)
    import traceback
    traceback.print_exc()
    sys.exit(1)
