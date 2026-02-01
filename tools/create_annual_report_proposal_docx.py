"""
年報変革提案書のWord文書生成スクリプト

HTML版の内容をpython-docxでWord文書として生成します。

Requirements:
- python-docx (already installed)

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_proposal_docx.py
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    print("python-docx imported successfully", flush=True)
    
    # 文書作成
    doc = Document()
    
    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'メイリオ'
    font.size = Pt(11)
    
    # ===== ページ1 =====
    
    # タイトル
    title = doc.add_heading('🌱 年報への変革 - 新しい自治会運営のために', level=1)
    title_format = title.paragraph_format
    title_format.space_after = Pt(10)
    
    # 提案の要点ボックス
    summary = doc.add_paragraph()
    summary_run = summary.add_run('📋 提案の要点\n')
    summary_run.bold = True
    summary_run.font.size = Pt(12)
    summary_text = summary.add_run('総会資料に加えて、「考え方と風土の共有」を目的とした年報を導入し、会員の理解と参加を促進する')
    summary_text.font.size = Pt(11)
    summary.paragraph_format.left_indent = Inches(0.3)
    summary.paragraph_format.space_after = Pt(10)
    # 背景色を設定（薄青）
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E8F4F8')
    summary._element.get_or_add_pPr().append(shading_elm)
    
    # セクション1
    section1_title = doc.add_heading('1. なぜ年報に変えるのか？', level=2)
    section1_title.paragraph_format.space_after = Pt(8)
    
    # 視覚的な問いかけボックス
    question_box = doc.add_paragraph()
    q_icon = question_box.add_run('❓ ')
    q_icon.font.size = Pt(14)
    q_text = question_box.add_run('こんなことありませんか？')
    q_text.bold = True
    q_text.font.size = Pt(10)
    question_box.paragraph_format.left_indent = Inches(0.3)
    question_box.paragraph_format.space_after = Pt(5)
    shading_q = OxmlElement('w:shd')
    shading_q.set(qn('w:fill'), 'E3F2FD')
    question_box._element.get_or_add_pPr().append(shading_q)
    
    q1 = doc.add_paragraph('• 総会資料を配っても、あまり読まれていない…', style='Normal')
    q1.paragraph_format.left_indent = Inches(0.5)
    q1.paragraph_format.space_after = Pt(2)
    q1.runs[0].font.size = Pt(9)
    
    q2 = doc.add_paragraph('• なぜこの決定になったのか、説明が足りない…', style='Normal')
    q2.paragraph_format.left_indent = Inches(0.5)
    q2.paragraph_format.space_after = Pt(8)
    q2.runs[0].font.size = Pt(9)
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(3)
    p1.add_run('従来：').bold = True
    p1.add_run('総会資料のみ（会則・報告・議案の通知）')
    
    p2 = doc.add_paragraph()
    p2.add_run('課題：').bold = True
    p2.paragraph_format.space_after = Pt(3)
    
    c1 = doc.add_paragraph('「なぜそうなったのか」が伝わらない', style='List Bullet')
    c1.paragraph_format.space_after = Pt(2)
    c1.runs[0].font.size = Pt(9)
    c2 = doc.add_paragraph('会員は受け身、対話の余地が少ない', style='List Bullet')
    c2.paragraph_format.space_after = Pt(2)
    c2.runs[0].font.size = Pt(9)
    c3 = doc.add_paragraph('考え方や風土が共有されず、関心低下', style='List Bullet')
    c3.paragraph_format.space_after = Pt(8)
    c3.runs[0].font.size = Pt(9)
    
    # セクション2
    section2_title = doc.add_heading('2. 総会資料 vs 年報', level=2)
    section2_title.paragraph_format.space_after = Pt(8)
    
    # 対比をより視覚的に
    compare_intro = doc.add_paragraph()
    compare_intro.add_run('📊 ').font.size = Pt(16)
    compare_intro.add_run('何が変わるの？').bold = True
    compare_intro.paragraph_format.space_after = Pt(8)
    
    # テーブルで横並びレイアウト
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # 左列：従来の総会資料
    left_cell = table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run('❌ 従来の総会資料')
    left_run.bold = True
    left_run.font.size = Pt(10)
    left_para.paragraph_format.space_after = Pt(5)
    # 背景色
    shading_left = OxmlElement('w:shd')
    shading_left.set(qn('w:fill'), 'F5F5F5')
    left_cell._element.get_or_add_tcPr().append(shading_left)
    
    # 箇条書き（簡潔に）
    items_left = ['会則・規則が先', '報告と決定の通知', '上から下への伝達']
    for item in items_left:
        p = left_cell.add_paragraph(f'• {item}', style='Normal')
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(9)
    
    # 右列：年報（新しい形）
    right_cell = table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_run = right_para.add_run('✅ 年報（新しい形）')
    right_run.bold = True
    right_run.font.size = Pt(10)
    right_para.paragraph_format.space_after = Pt(5)
    # 背景色
    shading_right = OxmlElement('w:shd')
    shading_right.set(qn('w:fill'), 'E8F9E8')
    right_cell._element.get_or_add_tcPr().append(shading_right)
    
    # 箇条書き（簡潔に）
    items_right = ['考え方の共有が先', '背景・理由の説明', '読んで楽しい']
    for item in items_right:
        p = right_cell.add_paragraph(f'• {item}', style='Normal')
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(9)
    
    # セクション3
    section3_title = doc.add_heading('3. 期待される効果', level=2)
    section3_title.paragraph_format.space_after = Pt(8)
    
    # 多様性の視覚化（簡略版）
    diversity_box = doc.add_paragraph()
    div_icon = diversity_box.add_run('👥 ')
    div_icon.font.size = Pt(14)
    div_text = diversity_box.add_run('会員の多様性を尊重')
    div_text.bold = True
    div_text.font.size = Pt(10)
    diversity_box.paragraph_format.space_after = Pt(5)
    
    # 記号で分布を表現（コンパクトに）
    distribution = doc.add_paragraph()
    distribution.add_run('低い ◀━━━━▶ 高い\n👤👤👤👤👤👤👤')
    distribution.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    distribution.paragraph_format.space_after = Pt(3)
    distribution.runs[0].font.size = Pt(9)
    
    diversity_note = doc.add_paragraph()
    diversity_note.add_run('（役員活動への許容力は人それぞれ）')
    diversity_note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    diversity_note.runs[0].font.size = Pt(8)
    diversity_note.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    diversity_note.paragraph_format.space_after = Pt(10)
    
    # 北風と太陽（コンパクトに）
    sun_box = doc.add_paragraph()
    sun_icon = sun_box.add_run('🌬️☀️ ')
    sun_icon.font.size = Pt(14)
    sun_text = sun_box.add_run('「北風と太陽」: 力ずくより、自らやりたくなる環境を')
    sun_text.font.size = Pt(9)
    sun_box.paragraph_format.left_indent = Inches(0.3)
    sun_box.paragraph_format.space_after = Pt(10)
    shading_sun = OxmlElement('w:shd')
    shading_sun.set(qn('w:fill'), 'FFFDE7')
    sun_box._element.get_or_add_pPr().append(shading_sun)
    
    # 効果の説明（簡潔に）
    effects_intro = doc.add_paragraph()
    effects_intro.add_run('✨ ').font.size = Pt(14)
    e_text = effects_intro.add_run('期待される効果：')
    e_text.bold = True
    e_text.font.size = Pt(10)
    effects_intro.paragraph_format.space_after = Pt(5)
    
    eff1 = doc.add_paragraph('• 理解の深まり - 会則の背景・理由が理解される', style='Normal')
    eff1.paragraph_format.left_indent = Inches(0.3)
    eff1.paragraph_format.space_after = Pt(3)
    eff1.runs[0].font.size = Pt(9)
    
    eff2 = doc.add_paragraph('• 共感の醸成 - 納得感が生まれる', style='Normal')
    eff2.paragraph_format.left_indent = Inches(0.3)
    eff2.paragraph_format.space_after = Pt(3)
    eff2.runs[0].font.size = Pt(9)
    
    eff3 = doc.add_paragraph('• 参加意欲 - 意見を言いたくなる', style='Normal')
    eff3.paragraph_format.left_indent = Inches(0.3)
    eff3.paragraph_format.space_after = Pt(3)
    eff3.runs[0].font.size = Pt(9)
    
    eff4 = doc.add_paragraph('• 風土形成 - 楽しい自治会へ', style='Normal')
    eff4.paragraph_format.left_indent = Inches(0.3)
    eff4.paragraph_format.space_after = Pt(8)
    eff4.runs[0].font.size = Pt(9)
    
    # ハイライトボックス
    highlight1 = doc.add_paragraph()
    h1_run1 = highlight1.add_run('💡 目指すのは：\n')
    h1_run1.bold = True
    h1_run2 = highlight1.add_run('会則や規則が先にあるのではなく、考え方の共有と相互理解が先にある自治会運営')
    h1_run2.bold = True
    highlight1.paragraph_format.left_indent = Inches(0.3)
    highlight1.paragraph_format.space_after = Pt(10)
    shading_h1 = OxmlElement('w:shd')
    shading_h1.set(qn('w:fill'), 'FFF9E6')
    highlight1._element.get_or_add_pPr().append(shading_h1)
    
    # 決裁事項
    decision = doc.add_paragraph()
    d_run1 = decision.add_run('🎯 今回の決裁事項：\n')
    d_run1.bold = True
    d_run2 = decision.add_run('次ページの構成での年報作成を承認いただけますか？（印刷発注：2月23日予定）')
    decision.paragraph_format.left_indent = Inches(0.3)
    decision.paragraph_format.space_after = Pt(15)
    shading_d = OxmlElement('w:shd')
    shading_d.set(qn('w:fill'), 'FFF9E6')
    decision._element.get_or_add_pPr().append(shading_d)
    
    # ===== 改ページ =====
    doc.add_page_break()
    
    # ===== ページ2 =====
    
    # タイトル
    title2 = doc.add_heading('📖 年報の構成（2026年度版）', level=1)
    
    intro = doc.add_paragraph('今年度から、総会資料に加えて「考え方の共有」を目的とした内容を組み込みます')
    intro.paragraph_format.space_after = Pt(15)
    
    # 第1部
    part1_title = doc.add_paragraph()
    part1_title_run = part1_title.add_run('第1部：考え方と風土の共有 ✨ 新規')
    part1_title_run.bold = True
    part1_title_run.font.size = Pt(13)
    part1_title_run.font.color.rgb = RGBColor(44, 95, 141)
    part1_title.paragraph_format.space_after = Pt(8)
    shading_p1t = OxmlElement('w:shd')
    shading_p1t.set(qn('w:fill'), '2C5F8D')
    part1_title._element.get_or_add_pPr().append(shading_p1t)
    part1_title_run.font.color.rgb = RGBColor(255, 255, 255)
    
    section1 = doc.add_paragraph()
    sec1_run = section1.add_run('📝 役員引受の悩み - 異なる視点を共有する')
    sec1_run.bold = True
    sec1_run.font.color.rgb = RGBColor(44, 95, 141)
    sec1_run.font.size = Pt(12)
    section1.paragraph_format.left_indent = Inches(0.3)
    section1.paragraph_format.space_after = Pt(8)
    
    content_intro = doc.add_paragraph()
    content_intro.add_run('📖 内容の例：')
    content_intro.paragraph_format.left_indent = Inches(0.5)
    content_intro.runs[0].bold = True
    content_intro.runs[0].font.size = Pt(10)
    content_intro.paragraph_format.space_after = Pt(5)
    
    doc.add_paragraph('役員引受が敬遠される理由（構造的問題）', style='List Bullet').runs[0].font.size = Pt(9)
    l1 = doc.paragraphs[-1]
    l1.paragraph_format.space_after = Pt(2)
    l1.paragraph_format.space_before = Pt(0)
    
    doc.add_paragraph('保守派・革新派、それぞれの前提', style='List Bullet').runs[0].font.size = Pt(9)
    l2 = doc.paragraphs[-1]
    l2.paragraph_format.space_after = Pt(2)
    l2.paragraph_format.space_before = Pt(0)
    
    doc.add_paragraph('「一律＝平等」を問い直す', style='List Bullet').runs[0].font.size = Pt(9)
    l3 = doc.paragraphs[-1]
    l3.paragraph_format.space_after = Pt(2)
    l3.paragraph_format.space_before = Pt(0)
    
    doc.add_paragraph('希望役員制という試み', style='List Bullet').runs[0].font.size = Pt(9)
    l4 = doc.paragraphs[-1]
    l4.paragraph_format.space_after = Pt(8)
    l4.paragraph_format.space_before = Pt(0)
    
    feature = doc.add_paragraph()
    f_run1 = feature.add_run('💡 特徴：')
    f_run1.bold = True
    f_run1.font.size = Pt(9)
    feature.add_run('決定事項の通知ではなく、考え方の前提を丁寧に説明し、対話を促す').font.size = Pt(9)
    feature.paragraph_format.left_indent = Inches(0.3)
    feature.paragraph_format.space_after = Pt(8)
    shading_f = OxmlElement('w:shd')
    shading_f.set(qn('w:fill'), 'E8F4F8')
    feature._element.get_or_add_pPr().append(shading_f)
    
    # 楽しく読めるポイントを追加
    fun_box = doc.add_paragraph()
    fun_icon = fun_box.add_run('🎨 ')
    fun_icon.font.size = Pt(14)
    fun_text = fun_box.add_run('楽しく読めるポイント：イラスト・図解、問いかけ、カラフルなレイアウト')
    fun_text.font.size = Pt(9)
    fun_box.paragraph_format.left_indent = Inches(0.3)
    fun_box.paragraph_format.space_after = Pt(8)
    shading_fun = OxmlElement('w:shd')
    shading_fun.set(qn('w:fill'), 'F3E5F5')
    fun_box._element.get_or_add_pPr().append(shading_fun)
    
    # 第2部
    part2_title = doc.add_paragraph()
    part2_title_run = part2_title.add_run('第2部：総会資料（従来通り）')
    part2_title_run.bold = True
    part2_title_run.font.size = Pt(13)
    part2_title_run.font.color.rgb = RGBColor(255, 255, 255)
    part2_title.paragraph_format.space_after = Pt(8)
    shading_p2t = OxmlElement('w:shd')
    shading_p2t.set(qn('w:fill'), '2C5F8D')
    part2_title._element.get_or_add_pPr().append(shading_p2t)
    
    # 総会議案
    agenda_title = doc.add_paragraph()
    at_icon = agenda_title.add_run('📋 ')
    at_icon.font.size = Pt(14)
    at_run = agenda_title.add_run('総会議案')
    at_run.bold = True
    at_run.font.color.rgb = RGBColor(44, 95, 141)
    agenda_title.paragraph_format.left_indent = Inches(0.3)
    agenda_title.paragraph_format.space_after = Pt(8)
    
    agendas = [
        '第1号議案：2025年度事業報告',
        '第2号議案：2025年度会計報告',
        '第3号議案：2026年度役員改選案',
        '第4号議案：2026年度事業計画案',
        '第5号議案：2026年度予算案',
        '第6号議案：会則の一部改正'
    ]
    
    for agenda in agendas:
        p = doc.add_paragraph(agenda)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(2)  # 明示的に2ptに設定
        p.paragraph_format.space_before = Pt(0)  # 前の間隔も0に
        p.paragraph_format.line_spacing = 1.0  # 行間も1.0に固定
    
    # 参考資料
    ref_title = doc.add_paragraph()
    ref_icon = ref_title.add_run('📚 ')
    ref_icon.font.size = Pt(14)
    rt_run = ref_title.add_run('参考資料')
    rt_run.bold = True
    rt_run.font.color.rgb = RGBColor(44, 95, 141)
    ref_title.paragraph_format.left_indent = Inches(0.3)
    ref_title.paragraph_format.space_before = Pt(10)
    ref_title.paragraph_format.space_after = Pt(8)
    
    refs = ['会則全体', '会員名簿', '会員住居地図']
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(2)  # 明示的に2ptに設定
        p.paragraph_format.space_before = Pt(0)  # 前の間隔も0に
        p.paragraph_format.line_spacing = 1.0  # 行間も1.0に固定
    
    # 今後の展開（コンパクトに）
    future = doc.add_paragraph()
    future.paragraph_format.space_before = Pt(10)
    fut_run1 = future.add_run('📌 今後の展開：試行 → 会員の反応を見て改善 → 楽しく読める年報へ')
    fut_run1.font.size = Pt(9)
    future.paragraph_format.left_indent = Inches(0.3)
    future.paragraph_format.space_after = Pt(10)
    shading_fut = OxmlElement('w:shd')
    shading_fut.set(qn('w:fill'), 'FFF9E6')
    future._element.get_or_add_pPr().append(shading_fut)
    
    # 理想のビジョン（コンパクトに）
    vision = doc.add_paragraph()
    vision.paragraph_format.space_before = Pt(10)
    vis_icon = vision.add_run('🌟 ')
    vis_icon.font.size = Pt(14)
    vis_text = vision.add_run('目指す自治会：💬気軽に意見 🤝理解し合う 😊楽しい 🌱一緒に育てる')
    vis_text.font.size = Pt(9)
    vision.paragraph_format.left_indent = Inches(0.3)
    vision.paragraph_format.space_after = Pt(8)
    shading_vis = OxmlElement('w:shd')
    shading_vis.set(qn('w:fill'), 'E8F5E9')
    vision._element.get_or_add_pPr().append(shading_vis)
    
    # フッター
    footer = doc.add_paragraph()
    footer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(15)
    footer_run = footer.add_run('2026年2月 第4回役員会資料\n事務局 高橋正剛')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Attribution
    attribution = doc.add_paragraph()
    attribution.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    attribution.paragraph_format.space_before = Pt(5)
    attr_run = attribution.add_run('Created with assistance from GitHub Copilot')
    attr_run.font.size = Pt(8)
    attr_run.font.italic = True
    attr_run.font.color.rgb = RGBColor(153, 153, 153)
    
    # 保存
    output_path = Path(__file__).parent.parent / 'docs' / 'community' / '2026__' / 'annual_report_reform_proposal.docx'
    doc.save(output_path)
    
    print(f"✓ SUCCESS! Word文書を作成しました", flush=True)
    print(f"✓ Output: {output_path.relative_to(Path.cwd())}", flush=True)
    
except ImportError as e:
    print(f"ERROR: Missing required library - {e}", file=sys.stderr, flush=True)
    print("\nInstall missing libraries:", flush=True)
    print('  & "C:\\Program Files\\Python313\\python.exe" -m pip install python-docx', flush=True)
    sys.exit(1)
    
except Exception as e:
    print(f"ERROR: {e}", file=sys.stderr, flush=True)
    import traceback
    traceback.print_exc()
    sys.exit(1)
