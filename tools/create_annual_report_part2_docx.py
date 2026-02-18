r"""
年報 第2部 Word文書生成スクリプト

annual_report_part2_data.md からデータを読み込んで Word ファイルとして出力します。

Requirements:
- python-docx (pip install python-docx)

Usage:
    cd C:\Users\takahashi\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_part2_docx.py
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_part2_docx.py --draft

Options:
    --draft    仮データ版であることを先頭ページに赤字注記として追加

Note:
    このスクリプトは第2部のみを生成します（第1部の生成は create_annual_report_part1_docx_v2.py）。
    合本が必要な場合は Word の「挿入 > ファイルからテキスト」または PDF 結合を使用してください。

⚠️ CONTENT SOURCE:
    docs/community/2026__/annual_report_part2_data.md
    内容を更新する場合はそのファイルを編集して、このスクリプトを再実行してください。
"""

import sys
import re
from pathlib import Path

# ============================================================
# CONFIGURATION
# ============================================================
_BASE       = Path(__file__).parent.parent
DATA_FILE   = _BASE / 'docs' / 'community' / '2026__' / 'annual_report_part2_data.md'
OUTPUT_FILE = _BASE / 'docs' / 'community' / '2026__' / 'annual_report_part2_v1.docx'
IS_DRAFT    = '--draft' in sys.argv


# ============================================================
# MARKDOWN PARSER
# ============================================================

def strip_comments(text: str) -> str:
    """HTMLコメント（<!-- ... -->）を除去する"""
    return re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)


def parse_inline(text: str) -> list:
    """
    **bold** テキストを認識して [(text, is_bold), ...] のリストを返す。
    マークダウンの太字のみを処理し、その他の書式は無視する。
    """
    result = []
    pos = 0
    for m in re.finditer(r'\*\*(.*?)\*\*', text):
        if m.start() > pos:
            result.append((text[pos:m.start()], False))
        result.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        result.append((text[pos:], False))
    return result or [(text, False)]


def parse_table(lines: list) -> dict | None:
    """
    マークダウンテーブル行をパースして dict を返す。
    separator 行（|---|---|）の前がヘッダー、後がデータ行。
    """
    headers = []
    rows    = []
    sep_seen = False

    for line in lines:
        s = line.strip()
        # セパレータ行: |---|---| または |:---|:---:|
        if re.match(r'^\|[-:| ]+\|$', s):
            sep_seen = True
            continue
        cells = [c.strip() for c in s.strip('|').split('|')]
        if not sep_seen:
            headers = cells
        else:
            rows.append(cells)

    return {'type': 'table', 'headers': headers, 'rows': rows} if headers else None


def load_blocks(path: Path) -> list:
    """
    マークダウンファイルをパースして構造化ブロックのリストを返す。

    ブロック型:
    - {'type': 'h2',     'content': str}           -- 大章見出し (## )
    - {'type': 'h3',     'content': str}           -- 中章見出し (### )
    - {'type': 'h4',     'content': str}           -- 小章見出し (#### )
    - {'type': 'para',   'content': str}           -- 通常段落
    - {'type': 'bullet', 'content': str}           -- 箇条書き (- )
    - {'type': 'table',  'headers': [...], 'rows': [[...]]}
    - {'type': 'hr'}                               -- 水平線 (---)
    """
    text  = strip_comments(path.read_text(encoding='utf-8'))
    lines = text.splitlines()
    blocks    = []
    i         = 0
    started   = False  # 最初の ## セクションが来るまでメタデータをスキップ

    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()

        # ── メタデータスキップ（最初の ## まで）──
        if not started:
            if s.startswith('## '):
                started = True
                # fall through to process this line
            else:
                i += 1
                continue

        # ── 空行 ──
        if not s:
            i += 1
            continue

        # ── 水平線 ──
        if s == '---':
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # ── テーブル ──
        if s.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            tbl = parse_table(tbl_lines)
            if tbl:
                blocks.append(tbl)
            continue

        # ── 見出し（長い prefix から順に判定）──
        matched_heading = False
        for prefix, htype in [('#### ', 'h4'), ('### ', 'h3'), ('## ', 'h2')]:
            if s.startswith(prefix):
                blocks.append({'type': htype, 'content': s[len(prefix):]})
                i += 1
                matched_heading = True
                break
        if matched_heading:
            continue

        # ── # タイトル行はスキップ ──
        if s.startswith('# '):
            i += 1
            continue

        # ── 箇条書き ──
        if s.startswith('- ') or s.startswith('* '):
            blocks.append({'type': 'bullet', 'content': s[2:]})
            i += 1
            continue

        # ── 通常段落 ──
        if s:
            blocks.append({'type': 'para', 'content': s})
        i += 1

    return blocks


# ============================================================
# WORD BUILDER
# ============================================================

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx が見つかりません", file=sys.stderr)
    print('Install: & "C:\\Program Files\\Python313\\python.exe" -m pip install python-docx')
    sys.exit(1)

# 色定数
NAVY  = RGBColor(0, 51, 102)
BLUE  = RGBColor(30, 80, 140)
GRAY  = RGBColor(102, 102, 102)
RED   = RGBColor(160, 0, 0)


def _add_page_number(para):
    """フッター段落にページ番号フィールドを追加する（Part1スクリプトと同じ方式）"""
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _cell_shade(cell, fill_hex: str):
    """テーブルセルに背景色を設定する"""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shd)


def _bottom_border(para, color_hex='1E5090', sz='12'):
    """段落に下罫線を追加する"""
    pPr  = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _inline_para(doc, text: str, size=11, bold=False, italic=False,
                 color=None, after=8, before=0, indent=0.0, align=None):
    """インライン太字（**...**）に対応した段落を追加して返す"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align:
        p.paragraph_format.alignment = align
    for txt, ib in parse_inline(text):
        r = p.add_run(txt)
        r.font.size = Pt(size)
        r.bold    = bold or ib
        r.italic  = italic
        if color:
            r.font.color.rgb = color
    return p


def _table_block(doc, headers: list, rows: list):
    """マークダウンテーブルを Word テーブルとして文書に追加する"""
    nc = len(headers)
    if nc == 0:
        return

    tbl = doc.add_table(rows=1 + len(rows), cols=nc)
    tbl.style = 'Table Grid'

    # ヘッダー行
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        _cell_shade(cell, 'D6E4F7')
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)

    # データ行
    for i, row in enumerate(rows):
        for j in range(nc):
            cell = tbl.rows[i + 1].cells[j]
            txt  = row[j] if j < len(row) else ''
            for t, ib in parse_inline(txt):
                r = cell.paragraphs[0].add_run(t)
                r.bold = ib
                r.font.size = Pt(10)

    # テーブル後スペーサー
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _render_block(doc, block: dict):
    """1ブロックを Word 文書に書き出す"""
    btype = block['type']

    if btype == 'h2':
        # 大章見出し：太字 14pt 濃紺 + 下罫線
        p = _inline_para(doc, block['content'], size=14, bold=True,
                         color=NAVY, before=16, after=8)
        _bottom_border(p)

    elif btype == 'h3':
        # 中章見出し：太字 12pt 中紺
        _inline_para(doc, block['content'], size=12, bold=True,
                     color=BLUE, before=10, after=6)

    elif btype == 'h4':
        # 小章見出し：太字 11pt 中紺
        _inline_para(doc, block['content'], size=11, bold=True,
                     color=BLUE, before=6, after=4)

    elif btype == 'para':
        _inline_para(doc, block['content'], size=11, after=8)

    elif btype == 'bullet':
        _inline_para(doc, '・' + block['content'], size=11, after=4, indent=0.3)

    elif btype == 'table':
        _table_block(doc, block['headers'], block['rows'])

    elif btype == 'hr':
        # 水平線：細い下罫線
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.space_before = Pt(6)
        _bottom_border(p, color_hex='AAAAAA', sz='6')


def build_doc(blocks: list, is_draft: bool) -> 'Document':
    """ブロックリストから Word 文書を構築して返す"""
    doc = Document()

    # ── ページ設定 (A4) ──
    sec = doc.sections[0]
    sec.page_height   = Inches(11.69)
    sec.page_width    = Inches(8.27)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)

    # ── デフォルトフォント ──
    doc.styles['Normal'].font.name = 'メイリオ'
    doc.styles['Normal'].font.size = Pt(11)

    # ── フッターにページ番号 ──
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(fp)

    # ── タイトルブロック ──
    doc.add_paragraph().paragraph_format.space_after = Pt(24)   # 上余白調整

    p = _inline_para(doc, '観音台第二自治会',
                     size=18, bold=True, color=NAVY, after=8,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    p = _inline_para(doc, '年報　第2部：総会資料',
                     size=24, bold=True, color=NAVY, after=6,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    p = _inline_para(doc, '2025年度定期総会　2026年3月28日（土）',
                     size=12, color=GRAY, after=20,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    if is_draft:
        p = _inline_para(doc, '【仮データ版：〔仮〕マークの数値は未確定。配布前に実数値への差し替えが必要】',
                         size=11, bold=True, color=RED, after=16,
                         align=WD_ALIGN_PARAGRAPH.CENTER)

    # タイトル下の区切り線
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after  = Pt(20)
    sep.paragraph_format.space_before = Pt(0)
    _bottom_border(sep, color_hex='1E5090', sz='16')

    # ── 本文コンテンツ ──
    for block in blocks:
        _render_block(doc, block)

    return doc


# ============================================================
# MAIN
# ============================================================

def main():
    print(f"Reading : {DATA_FILE.relative_to(_BASE)}", flush=True)
    if not DATA_FILE.exists():
        print(f"ERROR: データファイルが見つかりません: {DATA_FILE}", file=sys.stderr)
        sys.exit(1)

    blocks = load_blocks(DATA_FILE)
    print(f"Parsed  : {len(blocks)} blocks", flush=True)

    doc = build_doc(blocks, IS_DRAFT)
    doc.save(OUTPUT_FILE)

    print(f"✓ Output: {OUTPUT_FILE.relative_to(_BASE)}", flush=True)
    print(f"✓ Size  : {OUTPUT_FILE.stat().st_size:,} bytes", flush=True)
    if IS_DRAFT:
        print("  (--draft モード: 仮データ注記あり)", flush=True)
    else:
        print("  (通常モード)", flush=True)


if __name__ == '__main__':
    main()
