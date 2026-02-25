"""
書面協議書（役員会議決）Word文書生成スクリプト

入力: docs/community/2026__/written_board_meeting_draft.md
出力: docs/community/2026__/written_board_meeting.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent.parent
SRC_MD = BASE_DIR / "docs/community/2026__/written_board_meeting_draft.md"
OUT_DOCX = BASE_DIR / "docs/community/2026__/written_board_meeting.docx"

JP_FONT = "游明朝"
JP_FONT_HEADING = "游ゴシック"


def set_font(run, size_pt, bold=False, color=None, font_name=None):
    run.font.name = font_name or JP_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._r.rPr.rFonts.set(qn("w:eastAsia"), font_name or JP_FONT)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, 10.5)
    return p


def add_heading(doc, text, level, size_pt=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size_pt, bold=True, font_name=JP_FONT_HEADING)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_box_paragraph(doc, text, indent_cm=0.5):
    """回答欄・引用ブロック用：罫線付き段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.right_indent = Cm(indent_cm)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, 10.5)
    # 段落に罫線（上下）
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom", "left", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "888888")
        pBdr.append(bdr)
    pPr.append(pBdr)
    return p


def add_answer_box(doc, label):
    """回答欄ブロック"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    set_font(run, 10.5, bold=True, font_name=JP_FONT_HEADING)

    # 名前・役職行
    pn = doc.add_paragraph()
    pn.paragraph_format.left_indent = Cm(1)
    run2 = pn.add_run("氏名：_______________　　　役職：_______________")
    set_font(run2, 10.5)
    pn.paragraph_format.space_after = Pt(4)

    # 賛否チェック欄
    for sym, label_text in [("◎", "賛成"), ("×", "反対"), ("△", "保留　　ご意見：")]:
        pc = doc.add_paragraph()
        pc.paragraph_format.left_indent = Cm(2)
        pc.paragraph_format.space_before = Pt(0)
        pc.paragraph_format.space_after = Pt(0)
        run3 = pc.add_run(f"　{sym}　{label_text}")
        set_font(run3, 11)
    
    # 余白
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def set_page_margins(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)


def build_docx():
    doc = Document()
    set_page_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    # ===== タイトル =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run("書面協議書（役員会議決）")
    set_font(run, 16, bold=True, font_name=JP_FONT_HEADING)

    # ===== ヘッダー情報 =====
    for label, value in [
        ("発出日", "2026年2月23日"),
        ("発出者", "事務局長　高橋"),
        ("宛　先", "観音台1丁目第二自治会　役員各位"),
        ("回答締切", "2026年2月28日（土）"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}：")
        set_font(r1, 10.5, bold=True)
        r2 = p.add_run(value)
        set_font(r2, 10.5)

    doc.add_paragraph()

    # ===== 趣旨 =====
    add_heading(doc, "趣旨", 2, size_pt=11)

    paras = [
        "今年度の総会（3月28日）に向けて、下記2点について役員の皆さんのご意見をお聞きします。",
        "ひとつは、総会配布資料を三部構成の「年報」として発行する件です。もうひとつは、賛助会員制度の導入について、総会の場で会員全体に問う件です。",
        "総会直前のこの時期に改めてお諮りすることになり、ご迷惑をおかけします。お詫び申し上げます。",
    ]
    for text in paras:
        p = add_paragraph(doc, text)
        p.paragraph_format.space_after = Pt(6)

    p_last = add_paragraph(doc, "2月28日（土）までに、ひと言でもご回答いただけますと幸いです。")
    p_last.runs[0].font.bold = True
    p_last.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("─" * 40).paragraph_format.space_after = Pt(4)

    # ===== 議案1 =====
    add_heading(doc, "議案1：2025年度総会資料を「年報」として発行する件", 2, size_pt=12)

    add_heading(doc, "提案内容", 3, size_pt=10.5)
    add_paragraph(doc, "今年度の総会配布資料を、下記の三部構成で編集・発行します。")

    parts = [
        ("第1部（風土の共有と財政の見通し）— 事務局・高橋個人の問題提起として発行",
         "自治会の現状と今後の課題（三層モデルによる分析・財政見通し等）\n「役員会の公式見解」ではなく、「事務局・高橋個人の問題提起」として位置づけ、表紙にその旨を明記します。"),
        ("第2部（住民アンケート記録）— 同上",
         "夏季・冬季アンケートの結果と補足意見の記録"),
        ("第3部（総会資料）— 例年どおり",
         "事業報告・会計報告・役員改選案・事業計画案・予算案"),
    ]
    for bold_text, body_text in parts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(bold_text + "\n")
        set_font(r1, 10.5, bold=True)
        r2 = p.add_run(body_text)
        set_font(r2, 10.5)

    add_answer_box(doc, "【回答欄 — 議案1】")

    doc.add_paragraph("─" * 40).paragraph_format.space_after = Pt(4)

    # ===== 議案2 =====
    add_heading(doc, "議案2：「賛助会員」制度の導入方針を総会付議事項とする件", 2, size_pt=12)

    add_heading(doc, "提案内容", 3, size_pt=10.5)
    add_paragraph(doc, "下記の問いを今年度総会（3月28日）の正式議決事項として付議することを承認します。")

    add_box_paragraph(doc,
        "総会付議議案（案）\n"
        "「ゴミ集積所の利用のみを希望する世帯向けに、新しい会員区分として『賛助会員』制度を設けることの是非を問う。」"
    )

    p_note = doc.add_paragraph()
    p_note.paragraph_format.left_indent = Cm(0.5)
    r = p_note.add_run("賛助会員制度の内容そのものは総会が判断します。役員会として承認するのは、この問いを総会に諮るという手続きだけです。")
    set_font(r, 10)
    r.font.italic = True

    add_heading(doc, "背景", 3, size_pt=10.5)
    bullets = [
        "現在、ゴミ集積所利用のみを希望する世帯には12,000円/年の非会員負担のみという状況",
        "自治会への期待が多様化する中、正会員区分の一本化が実態と乖離しつつある",
        "今年度実施したアンケートで「多様化許容」を支持する回答が多数寄せられている",
        "制度の詳細（会費額・会則改正・実施時期）は、総会の承認後に次年度役員会で確定する",
    ]
    for b in bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"・{b}")
        set_font(r, 10.5)

    add_answer_box(doc, "【回答欄 — 議案2】")

    doc.add_paragraph("─" * 40).paragraph_format.space_after = Pt(4)

    # ===== 回答方法 =====
    add_heading(doc, "回答方法", 2, size_pt=11)

    add_paragraph(doc, "2月28日（土）までに、以下のいずれの方法でも構いません。お手軽な方法でご回答ください。")
    for i, item in enumerate([
        "この書面に記入の上、事務局（高橋）まで手渡し・投函",
        "メール・電話・口頭など、連絡しやすい手段で「議案1は賛成／議案2は賛成」のようにお伝えいただくだけで結構です",
    ], 1):
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.75)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(3)
        r = p2.add_run(f"{i}．{item}")
        set_font(r, 10.5)

    p3 = add_paragraph(doc, "回答がない場合は「委任」として扱い、委任者の票は出席者の多数決に従います。")
    p3.runs[0].font.italic = True

    doc.add_paragraph("─" * 40).paragraph_format.space_after = Pt(4)

    # ===== 補足：手続きの根拠 =====
    add_heading(doc, "補足：手続きの根拠", 2, size_pt=11)

    add_paragraph(doc,
        "本書面協議は、会則第25条「その他、本会の運営に関し必要な事項は、その都度、役員会で協議して決定する」に基づく"
        "正式な役員会の議決として行います。会長を含む各役員は、賛成・反対のいずれかを表明する権利と責任を持ちます。"
    )

    for bold_text, body_text in [
        ("定足数", "役員9名中6名以上の回答"),
        ("議決", "回答者の過半数の賛成"),
        ("回答なしの場合", "委任として扱い、出席者の多数決に従います"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"・{bold_text}：")
        set_font(r1, 10.5, bold=True)
        r2 = p.add_run(body_text)
        set_font(r2, 10.5)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"✓ 生成完了: {OUT_DOCX}", flush=True)
    print(f"  サイズ: {OUT_DOCX.stat().st_size:,} bytes", flush=True)


if __name__ == "__main__":
    build_docx()
