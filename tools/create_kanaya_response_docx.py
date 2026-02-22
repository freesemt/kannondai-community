r"""
金谷氏への中間応答レター 生成スクリプト

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_kanaya_response_docx.py

データソース:
    docs/community/2026__/kanaya_interim_response_draft.md

出力:
    docs/community/2026__/kanaya_interim_response.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = Path(__file__).parent.parent
DRAFT  = BASE / 'docs/community/2026__/kanaya_interim_response_draft.md'
OUTPUT = BASE / 'docs/community/2026__/kanaya_interim_response.docx'

NAVY      = RGBColor(0x17, 0x37, 0x5E)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)

# ============================================================
# ユーティリティ
# ============================================================

def set_default_font(doc, font_name='游明朝', size_pt=11):
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def set_run_font(run, font_name='游明朝'):
    run.font.name = font_name
    run._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), font_name)


def add_para(doc, text, font_size=11, indent_left=0,
             space_before=None, space_after=6,
             color=None, bold=False, align=None):
    """インライン ** を太字として処理しながら段落を追加する"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if indent_left:
        pf.left_indent = Inches(indent_left)
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # インライン ** 処理
    segments = re.split(r'(\*\*.*?\*\*)', text)
    for seg in segments:
        m = re.match(r'\*\*(.*?)\*\*', seg)
        run = p.add_run(m.group(1) if m else seg)
        run.bold = bool(m) or bold
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
        set_run_font(run)
    return p


def add_heading_line(doc, text, font_size=12, color=None):
    """小見出し行（左罫線なし、シンプル太字）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = color or NAVY
    set_run_font(run)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('―' * 30)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    set_run_font(run)


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.2):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


# ============================================================
# メイン
# ============================================================

def main():
    raw = DRAFT.read_text(encoding='utf-8')

    # ## 本文 の下だけを使う
    body_match = re.search(r'^## 本文\s*\n(.+)', raw, re.DOTALL | re.MULTILINE)
    if not body_match:
        print('[ERROR] "## 本文" セクションが見つかりません')
        return
    body = body_match.group(1).strip()

    doc = Document()
    set_default_font(doc)
    set_page_margins(doc)

    count = 0

    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        # --- セクション区切り ---
        if stripped == '---':
            add_separator(doc)

        # --- 小見出し（**...** のみの行 または --- で囲まれた太字行）---
        elif re.match(r'^\*\*.+\*\*$', stripped):
            text = re.sub(r'^\*\*|\*\*$', '', stripped)
            add_heading_line(doc, text, font_size=11, color=NAVY)

        # --- 宛名（「金谷 様」）---
        elif stripped == '金谷 様':
            add_para(doc, stripped, font_size=12, bold=True, space_before=6, space_after=12)

        # --- 敬具 ---
        elif stripped == '敬具':
            add_para(doc, stripped, font_size=11, align='right', space_before=10, space_after=4)

        # --- 日付・署名行 ---
        elif stripped.startswith('令和') or stripped.startswith('観音台') or stripped.startswith('　高橋'):
            add_para(doc, stripped, font_size=11, align='right', space_after=4)

        # --- 通常段落 ---
        else:
            add_para(doc, stripped, font_size=11, space_after=7)

        count += 1

    doc.save(OUTPUT)
    print(f'[OK] 出力完了: {OUTPUT}  ({count} 要素)')


if __name__ == '__main__':
    main()
