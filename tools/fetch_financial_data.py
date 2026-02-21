"""
過去7年分の総会資料から会計データを抽出するスクリプト

対象ファイル（tools/documents/）:
  - 総会資料平成30年度 P1_20 表紙と事業報告.docx  (H30, FY2018)
  - 令和元年度総会資料_200302.docx                  (R1,  FY2019)
  - 令和2年度総会資料.docx                          (R2,  FY2020)
  - 令和3年度総会資料.docx                          (R3,  FY2021)
  - 令和4年度総会資料.docx                          (R4,  FY2022)
  - 令和5年度総会資料0310.docx                      (R5,  FY2023)
  - 2024年度総会資料0330最終版.docx                 (2024年度)

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\fetch_financial_data.py
"""

import sys
from pathlib import Path
from docx import Document


# 対象ファイルと年度のマッピング
TARGET_FILES = [
    ("総会資料平成30年度 P1_20 表紙と事業報告.docx", "平成30年度 (2018)"),
    ("令和元年度総会資料_200302.docx",               "令和元年度 (2019)"),
    ("令和2年度総会資料.docx",                       "令和2年度 (2020)"),
    ("令和3年度総会資料.docx",                       "令和3年度 (2021)"),
    ("令和4年度総会資料.docx",                       "令和4年度 (2022)"),
    ("令和5年度総会資料0310.docx",                   "令和5年度 (2023)"),
    ("2024年度総会資料0330最終版.docx",               "2024年度 (2024)"),
]

SCRIPT_DIR = Path(__file__).parent
DOCUMENTS_DIR = SCRIPT_DIR / "documents"
OUTPUT_FILE = SCRIPT_DIR / "financial_data_extracted.txt"


def extract_tables(doc_path):
    """docxファイルのすべてのテーブルをリスト形式で返す"""
    doc = Document(doc_path)
    tables_data = []
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # 完全に空の行はスキップ
            if any(c for c in cells):
                rows.append(cells)
        if rows:
            tables_data.append(rows)
    return tables_data


def extract_paragraphs_near_accounting(doc_path):
    """会計報告セクション周辺の段落テキストを返す"""
    doc = Document(doc_path)
    lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return lines


def format_table(rows):
    """テーブル行をテキスト形式に整形"""
    lines = []
    for row in rows:
        lines.append("  | " + " | ".join(row) + " |")
    return "\n".join(lines)


def main():
    print("=" * 60, flush=True)
    print("会計データ抽出スクリプト", flush=True)
    print("=" * 60, flush=True)

    all_output = []
    all_output.append("# 過去7年分 会計データ抽出結果\n")
    all_output.append(f"抽出元: {DOCUMENTS_DIR}\n")
    all_output.append("=" * 60 + "\n\n")

    for filename, label in TARGET_FILES:
        doc_path = DOCUMENTS_DIR / filename
        print(f"\n--- {label} ---", flush=True)

        if not doc_path.exists():
            msg = f"  ⚠ ファイルなし: {filename}"
            print(msg, flush=True)
            all_output.append(f"\n## {label}\n{msg}\n")
            continue

        try:
            # テーブル抽出
            tables = extract_tables(doc_path)
            # 段落（見出し用）
            paragraphs = extract_paragraphs_near_accounting(doc_path)

            print(f"  テーブル数: {len(tables)}", flush=True)
            print(f"  段落数: {len(paragraphs)}", flush=True)

            section_lines = [f"\n## {label}\n"]

            # 会計関連の段落抽出（「会計報告」「収入」「支出」「繰越」を含む行）
            accounting_keywords = ["会計報告", "収入の部", "支出の部", "繰越の部",
                                   "一般会計", "特別会計", "積立金", "委託料",
                                   "会費", "報告します", "会計監査", "予備費"]
            relevant_paras = [p for p in paragraphs
                              if any(kw in p for kw in accounting_keywords)]
            if relevant_paras:
                section_lines.append("### 会計関連段落\n")
                for p in relevant_paras:
                    section_lines.append(f"  {p}")
                section_lines.append("")

            # テーブル全件出力（番号付き）
            if tables:
                section_lines.append("### テーブルデータ\n")
                for i, table_rows in enumerate(tables, 1):
                    # テーブルが会計っぽいか判定（数字・円を含むか）
                    flat = " ".join(" ".join(r) for r in table_rows)
                    is_financial = any(c in flat for c in ["円", "￥", "収入", "支出", "繰越", "委託", "積立"])
                    marker = "💰" if is_financial else "  "
                    section_lines.append(f"  {marker} テーブル {i}:")
                    section_lines.append(format_table(table_rows))
                    section_lines.append("")
            else:
                section_lines.append("  （テーブルなし）\n")

            all_output.extend(section_lines)

        except Exception as e:
            msg = f"  ❌ エラー: {e}"
            print(msg, flush=True)
            all_output.append(f"\n## {label}\n{msg}\n")

    # ファイル出力
    output_text = "\n".join(all_output)
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        f.write(output_text)

    print(f"\n✓ 出力完了: {OUTPUT_FILE.relative_to(Path.cwd())}", flush=True)
    print(f"  合計 {len(output_text)} 文字", flush=True)


if __name__ == "__main__":
    main()
