"""
年報 第1部（前半＋後半）見本 生成スクリプト

前半: annual_report_part1_draft_v2.md  (セクション1〜5)
後半: annual_report_part1_draft_v2_second.md (セクション6〜10)
グラフ: tools/forecast_garbage_cumulative.png

出力: docs/community/2026__/annual_report_part1_v2_mitsumori.docx

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_part1_v2_mitsumori.py
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", flush=True)
    sys.exit(1)

BASE = Path(__file__).parent.parent
DOCS = BASE / 'docs' / 'community' / '2026__'
TOOLS = BASE / 'tools'

MD_FIRST  = DOCS / 'annual_report_part1_draft_v2.md'
MD_SECOND = DOCS / 'annual_report_part1_draft_v2_second.md'
GRAPH     = TOOLS / 'forecast_garbage_cumulative.png'
HIER_IMG  = DOCS / 'images' / 'hierarchy_diagram.png'
OUTPUT    = DOCS / 'annual_report_part1_v2_mitsumori.docx'


def extract_body(md_path):
    """## ドラフト本文 〜 ## 執筆メモ の間を抽出して構造リストに変換"""
    text = md_path.read_text(encoding='utf-8')
    m = re.search(r'## ドラフト本文.*?\n\n(.*?)(?=\n## 執筆メモ|\Z)', text, re.DOTALL)
    if not m:
        raise ValueError(f"ドラフト本文が見つかりません: {md_path}")

    lines = m.group(1).split('\n')
    items = []
    buf = []

    def flush_buf():
        if buf:
            items.append({'type': 'para', 'content': '\n'.join(buf)})
            buf.clear()

    for line in lines:
        stripped = line.strip()

        if stripped == '---':
            flush_buf()
        elif stripped.startswith('### '):
            flush_buf()
            items.append({'type': 'heading', 'content': stripped[4:]})
        elif stripped.startswith('| '):
            # Markdown table row
            flush_buf()
            items.append({'type': 'table_row', 'content': stripped})
        elif stripped.startswith('> '):
            flush_buf()
            items.append({'type': 'quote', 'content': stripped[2:]})
        elif stripped.startswith('!['):
            flush_buf()
            img_m = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
            if img_m:
                items.append({'type': 'image', 'alt': img_m.group(1), 'src': img_m.group(2)})
        elif stripped.startswith('**[図：'):
            flush_buf()
            items.append({'type': 'image_placeholder', 'content': stripped})
        elif stripped:
            buf.append(stripped)
        else:
            flush_buf()

    flush_buf()

    # 隣接する table_row をまとめる
    merged = []
    i = 0
    while i < len(items):
        if items[i]['type'] == 'table_row':
            rows = []
            while i < len(items) and items[i]['type'] == 'table_row':
                rows.append(items[i]['content'])
                i += 1
            merged.append({'type': 'table', 'rows': rows})
        else:
            merged.append(items[i])
            i += 1
    return merged


def parse_bold(text):
    """**bold** を (text, is_bold) リストに分解"""
    parts = []
    for m in re.finditer(r'(\*\*.*?\*\*|[^*]+|\*)', text):
        t = m.group(0)
        if t.startswith('**') and t.endswith('**'):
            parts.append((t[2:-2], True))
        else:
            parts.append((t, False))
    return parts or [(text, False)]


def add_para(doc, text, indent=0, size=11, space_after=10, line_spacing=1.3):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for t, bold in parse_bold(text):
        r = p.add_run(t)
        r.bold = bold
        r.font.size = Pt(size)
    return p


def add_image(doc, img_path, caption_text):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(str(img_path), width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(14)
        cr = cap.add_run(caption_text)
        cr.font.size = Pt(9)
        cr.italic = True
        cr.font.color.rgb = RGBColor(102, 102, 102)
    else:
        print(f"  ⚠ 画像が見つかりません: {img_path}", flush=True)


def add_table_from_rows(doc, rows):
    """Markdownテーブル行リストをWordテーブルに変換（区切り行は除外）"""
    data = []
    for row in rows:
        if re.match(r'\|[\s\-:|]+\|', row):
            continue   # 区切り行をスキップ
        cells = [c.strip() for c in row.strip('|').split('|')]
        data.append(cells)
    if not data:
        return
    ncols = max(len(r) for r in data)
    tbl = doc.add_table(rows=len(data), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(data):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            c = tbl.rows[ri].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            for t, bold in parse_bold(cell_text):
                r = p.add_run(t)
                r.bold = bold
                r.font.size = Pt(10)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


def render_sections(doc, sections, graph_path, hier_path):
    """構造リストをWordに書き出す"""
    inserted_hier = False
    inserted_graph = False

    for item in sections:
        t = item['type']

        if t == 'heading':
            h = doc.add_heading(item['content'], level=3)
            h.paragraph_format.space_after = Pt(8)
            h.paragraph_format.left_indent = Inches(0)
            h.paragraph_format.first_line_indent = Inches(0)

        elif t == 'para':
            content = item['content']
            # 「[図：...」マーカーがあれば階層図を挿入（セクション2用）
            if '[図：' in content and not inserted_hier:
                add_image(doc, hier_path, '図：自治会への期待：時代による変化')
                inserted_hier = True
                continue
            add_para(doc, content)

        elif t == 'image':
            src = item['src']
            # forecast_garbage_cumulative.png → GRAPH を使用
            if 'forecast_garbage' in src:
                if not inserted_graph:
                    add_image(doc, graph_path,
                              '図：ゴミ集積所使用料 累計推移（上：総資産推移 ／ 下：非会員数・使用料累計）')
                    inserted_graph = True
            elif 'hierarchy_diagram' in src:
                if not inserted_hier:
                    add_image(doc, hier_path, '図：自治会への期待：時代による変化')
                    inserted_hier = True

        elif t == 'image_placeholder':
            if not inserted_graph:
                add_image(doc, graph_path,
                          '図：ゴミ集積所使用料 累計推移（上：総資産推移 ／ 下：非会員数・使用料累計）')
                inserted_graph = True

        elif t == 'table':
            add_table_from_rows(doc, item['rows'])

        elif t == 'quote':
            add_para(doc, item['content'], indent=0.4, size=10, space_after=6)


# ============================================================
# 文書生成
# ============================================================
print("年報 第1部 見本v2 生成開始", flush=True)

sections_first  = extract_body(MD_FIRST)
sections_second = extract_body(MD_SECOND)
print(f"  前半: {len(sections_first)} ブロック", flush=True)
print(f"  後半: {len(sections_second)} ブロック", flush=True)

doc = Document()

# ページ設定（A4）
sec = doc.sections[0]
sec.page_height = Inches(11.69)
sec.page_width  = Inches(8.27)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin   = Inches(1.1)
sec.right_margin  = Inches(1.1)

# デフォルトフォント
style = doc.styles['Normal']
style.font.name = 'メイリオ'
style.font.size = Pt(11)

# ----- 表紙 -----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(60)
title_p.paragraph_format.space_after  = Pt(10)
tr = title_p.add_run('観音台第二自治会　年報 2025（見本）')
tr.font.size = Pt(26)
tr.bold = True
tr.font.color.rgb = RGBColor(0, 51, 102)

note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_p.paragraph_format.space_after = Pt(60)
nr = note_p.add_run('― 印刷費確認用見本 ―')
nr.font.size = Pt(13)
nr.italic = True
nr.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# ----- 第1部：前半 -----
h1 = doc.add_heading('第1部：考え方と風土の共有', level=2)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_after = Pt(18)

render_sections(doc, sections_first, GRAPH, HIER_IMG)

# ----- 第1部：後半（財政データ）-----
doc.add_page_break()

h2 = doc.add_heading('第1部（続き）：数字で見る30年の見通し', level=2)
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h2.paragraph_format.space_after = Pt(18)

render_sections(doc, sections_second, GRAPH, HIER_IMG)

# ----- 末尾注記 -----
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(30)
note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
nr2 = note.add_run('2026年2月　観音台自治会 事務局（見本）')
nr2.font.size = Pt(9)
nr2.font.color.rgb = RGBColor(128, 128, 128)

# ----- 保存 -----
doc.save(str(OUTPUT))
print(f"\n✅ 保存完了: {OUTPUT}", flush=True)
