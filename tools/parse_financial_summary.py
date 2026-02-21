"""
7年分の財務データを解析して年度比較表を生成するスクリプト
fetch_financial_data.py で生成した financial_data_extracted.txt を元に、
主要な会計指標（収入・支出・繰越）を年度別にまとめた Markdown 表を出力する。

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\parse_financial_summary.py
"""

import re
from pathlib import Path
from docx import Document


# 対象ファイルと年度のマッピング
TARGET_FILES = [
    ("総会資料平成30年度 P1_20 表紙と事業報告.docx", "平成30年度", "2018"),
    ("令和元年度総会資料_200302.docx",               "令和元年度", "2019"),
    ("令和2年度総会資料.docx",                       "令和2年度",  "2020"),
    ("令和3年度総会資料.docx",                       "令和3年度",  "2021"),
    ("令和4年度総会資料.docx",                       "令和4年度",  "2022"),
    ("令和5年度総会資料0310.docx",                   "令和5年度",  "2023"),
    ("2024年度総会資料0330最終版.docx",               "2024年度",   "2024"),
]

SCRIPT_DIR = Path(__file__).parent
DOCUMENTS_DIR = SCRIPT_DIR / "documents"
OUTPUT_MD   = SCRIPT_DIR / "financial_summary.md"
OUTPUT_CSV  = SCRIPT_DIR / "financial_summary.csv"


def to_int(s):
    """コンマ区切り・▲（マイナス）を含む数値文字列を int に変換。失敗時は None"""
    s = s.strip().replace(",", "").replace("，", "")
    if not s:
        return None
    negative = s.startswith("▲")
    s = s.lstrip("▲")
    try:
        v = int(s)
        return -v if negative else v
    except ValueError:
        return None


def parse_financials(doc_path):
    """
    docx からメイン会計テーブルを探し、
    {
      "一般収入": int,
      "一般支出": int,
      "一般前繰越": int,
      "一般次繰越": int,
      "世帯数": int,
      "委託料": int,
      "自治会費収入": int,
    }
    を返す。見つからない場合は None。
    """
    doc = Document(doc_path)
    result = {}

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        flat = " ".join(" ".join(r) for r in rows)

        # 一般会計の収支表を識別
        if "収入の部" not in flat and "支出の部" not in flat:
            continue
        if "合　計" not in flat and "合計" not in flat:
            continue

        # --- 収入合計 ---
        for row in rows:
            if any(c in ["合　計", "合計"] for c in row):
                # 収入の部の合計行：「自治会費」行の後かどうかを確認するより
                # 最初に出てくる「収入の部」直後の合計を使う
                nums = [to_int(c) for c in row if to_int(c) is not None and to_int(c) > 50000]
                if nums and "一般収入" not in result:
                    result["一般収入"] = nums[0]

        # --- 支出合計 ---
        found_shunyu = False
        for row in rows:
            joined = " ".join(row)
            if "収入の部" in joined:
                found_shunyu = True
            if "支出の部" in joined:
                found_shunyu = False  # 支出セクション開始
            if not found_shunyu and any(c in ["合　計", "合計"] for c in row):
                nums = [to_int(c) for c in row if to_int(c) is not None and to_int(c) > 50000]
                if nums and "一般支出" not in result:
                    result["一般支出"] = nums[0]

        # --- 繰越金 ---
        for row in rows:
            joined = " ".join(row)
            # 前年度繰越
            if ("前年度繰越" in joined or "前年繰越" in joined) and "一般前繰越" not in result:
                nums = [to_int(c) for c in row if to_int(c) is not None and to_int(c) > 0]
                if nums:
                    result["一般前繰越"] = nums[0]
            # 次年度繰越
            if ("次年度繰越" in joined or "次年繰越" in joined) and "一般次繰越" not in result:
                nums = [to_int(c) for c in row if to_int(c) is not None and to_int(c) > 0]
                if nums:
                    result["一般次繰越"] = nums[0]

        # --- 委託料 ---
        for row in rows:
            joined = " ".join(row)
            if ("委託料" in joined or "委託金" in joined) and "委託料" not in result:
                nums = [to_int(c) for c in row if to_int(c) is not None and to_int(c) > 50000]
                if nums:
                    result["委託料"] = nums[0]

        # --- 自治会費収入 ---
        for row in rows:
            joined = " ".join(row)
            if "自治会費" in joined and "一般" not in joined and "自治会費収入" not in result:
                nums = [to_int(c) for c in row if to_int(c) is not None and 50000 < to_int(c) < 1000000]
                if nums:
                    result["自治会費収入"] = nums[0]

        if result:
            break  # 最初の会計テーブルで十分

    # --- 世帯数 ---
    for para in doc.paragraphs:
        m = re.search(r"(\d+)\s*[世戸]", para.text)
        if m:
            n = int(m.group(1))
            if 80 < n < 200:
                result.setdefault("世帯数", n)

    return result if result else None


def fmt(v, unit="円"):
    if v is None:
        return "–"
    if unit == "円":
        return f"{v:,}"
    return str(v)


def main():
    print("=" * 60)
    print("財務サマリー生成スクリプト")
    print("=" * 60)

    records = []
    for filename, label, year in TARGET_FILES:
        doc_path = DOCUMENTS_DIR / filename
        print(f"\n{label} ({year}) ...", end="", flush=True)
        if not doc_path.exists():
            print(" ファイルなし")
            records.append((label, year, None))
            continue
        data = parse_financials(doc_path)
        if data:
            print(f" OK  収入={fmt(data.get('一般収入'))} / 支出={fmt(data.get('一般支出'))} / 次繰越={fmt(data.get('一般次繰越'))}")
        else:
            print(" データ取得失敗")
        records.append((label, year, data))

    # --- Markdown 出力 ---
    lines = []
    lines.append("# 観音台第二自治会 過去7年分 会計サマリー\n")
    lines.append("（一般会計 決算ベース、単位：円）\n")

    lines.append("## 年度別収支一覧\n")
    lines.append("| 年度 | 世帯数 | 収入合計 | うち会費 | うち委託料 | 支出合計 | 前繰越 | 次繰越 | 収支差額 |")
    lines.append("|------|--------|----------|----------|------------|----------|--------|--------|----------|")

    for label, year, data in records:
        if data is None:
            lines.append(f"| {label} | – | – | – | – | – | – | – | – |")
            continue
        収入 = data.get("一般収入")
        支出 = data.get("一般支出")
        前繰 = data.get("一般前繰越")
        次繰 = data.get("一般次繰越")
        会費 = data.get("自治会費収入")
        委託 = data.get("委託料")
        世帯 = data.get("世帯数")
        差額 = (収入 - 支出) if 収入 is not None and 支出 is not None else None
        lines.append(
            f"| {label} | {fmt(世帯,'')} | {fmt(収入)} | {fmt(会費)} | {fmt(委託)} "
            f"| {fmt(支出)} | {fmt(前繰)} | {fmt(次繰)} | {fmt(差額)} |"
        )

    lines.append("")
    lines.append("## 繰越金推移\n")
    lines.append("| 年度 | 次年度繰越（一般） |")
    lines.append("|------|------------------|")
    for label, year, data in records:
        if data is None:
            lines.append(f"| {label} | – |")
        else:
            lines.append(f"| {label} | {fmt(data.get('一般次繰越'))} |")

    lines.append("")
    lines.append("> 注: 上記数値は総会資料の一般会計決算テーブルから自動抽出したものです。")
    lines.append("> 特別会計（集会所積立金）は別途参照してください。")

    md_text = "\n".join(lines)
    with open(OUTPUT_MD, "w", encoding="utf-8") as f:
        f.write(md_text)
    print(f"\n✓ Markdown 出力: {OUTPUT_MD.relative_to(Path.cwd())}")

    # --- CSV 出力 ---
    csv_lines = ["年度,西暦,世帯数,収入合計,うち会費,うち委託料,支出合計,収支差額,前年繰越,次年繰越"]
    for label, year, data in records:
        if data is None:
            csv_lines.append(f"{label},{year},,,,,,,,")
            continue
        def v(k): return data.get(k, "")
        収入 = data.get("一般収入", "")
        支出 = data.get("一般支出", "")
        差額 = (収入 - 支出) if isinstance(収入, int) and isinstance(支出, int) else ""
        csv_lines.append(
            f"{label},{year},{v('世帯数')},{収入},{v('自治会費収入')},{v('委託料')},"
            f"{支出},{差額},{v('一般前繰越')},{v('一般次繰越')}"
        )
    with open(OUTPUT_CSV, "w", encoding="utf-8") as f:
        f.write("\n".join(csv_lines))
    print(f"✓ CSV 出力:      {OUTPUT_CSV.relative_to(Path.cwd())}")


if __name__ == "__main__":
    main()
