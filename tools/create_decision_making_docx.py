#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
意思決定プロセス改善文書（Word形式）生成スクリプト
2ページ以内でコンパクトに
"""

import os
import sys

# python-docxのインポート
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    print("python-docx imported successfully")
except ImportError:
    print("ERROR: python-docx is not installed")
    print("Please run: pip install python-docx")
    sys.exit(1)

def add_colored_box(paragraph, text, bg_color):
    """背景色付きボックスを追加"""
    run = paragraph.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    
    # 背景色設定
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), bg_color)
    run._element.get_or_add_rPr().append(shading_elm)
    
    return run

def create_decision_making_document():
    """意思決定プロセス改善Word文書を作成"""
    
    doc = Document()
    
    # ページ設定（マージンを狭く）
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # ========== タイトル ==========
    title = doc.add_heading('どうやって決めていますか？', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('自治会の意思決定を考える')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(6)
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    # ========== エピソード ==========
    p1 = doc.add_paragraph()
    p1.add_run('🏞️ ある日の役員会で').font.size = Pt(12)
    p1.runs[0].font.bold = True
    p1.paragraph_format.space_after = Pt(3)
    
    p2 = doc.add_paragraph('羽成公園の遊具を新しくする提案が届きました。役員会で資料を回覧し、')
    p2.add_run('その場で全員一致で決定').font.bold = True
    p2.add_run('しました。')
    for run in p2.runs:
        run.font.size = Pt(10)
    p2.paragraph_format.space_after = Pt(2)
    
    p3 = doc.add_paragraph()
    p3.add_run('💭 でも後から「').font.size = Pt(10)
    run = p3.add_run('十分に検討できたのだろうか？')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    p3.add_run('」という疑問が残りました。').font.size = Pt(10)
    p3.paragraph_format.space_after = Pt(8)
    
    # ========== 2つの決め方の対比 ==========
    heading2 = doc.add_heading('📊 2つの決め方を比べてみる', level=2)
    heading2.runs[0].font.size = Pt(13)
    heading2.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading2.paragraph_format.space_after = Pt(4)
    
    # 表を作成（2列）
    table = doc.add_table(rows=8, cols=2)  # ヘッダー1 + データ7 = 8行
    table.style = 'Light Grid Accent 1'
    
    # ヘッダー行
    header_cells = table.rows[0].cells
    header_cells[0].text = '❌ 良くない例：遊具選定'
    header_cells[1].text = '✅ 改善した例：今年度総会'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E6E6FA')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # データ行
    row_data = [
        ('提案届く', '3/1 年報配布'),
        ('↓', '↓ （4週間の検討期間）'),
        ('その場で回覧', '3/7 質疑応接（1回目）'),
        ('↓ （数分）', '3/14 質疑応接（2回目）'),
        ('即決定', '↓'),
        ('', '3/21 意見回収'),
        ('', '3/28 総会で決定'),
    ]
    
    for i, (left, right) in enumerate(row_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            if cell.text:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    
    # 表の後のスペース
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # ========== 対比表 ==========
    comparison_table = doc.add_table(rows=5, cols=3)
    comparison_table.style = 'Light List Accent 1'
    
    # ヘッダー
    hdr_cells = comparison_table.rows[0].cells
    hdr_cells[0].text = '項目'
    hdr_cells[1].text = '遊具選定'
    hdr_cells[2].text = '総会プロセス'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    # データ
    comparison_data = [
        ('事前情報', 'その場で見る', '4週間前に配布'),
        ('考える時間', '数分', '4週間'),
        ('質問の機会', 'その場のみ', '2回＋意見用紙'),
        ('納得感', '？', '高い'),
    ]
    
    for i, (item, bad, good) in enumerate(comparison_data, start=1):
        cells = comparison_table.rows[i].cells
        cells[0].text = item
        cells[1].text = bad
        cells[2].text = good
        for cell in cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # ========== 改善ポイント ==========
    heading3 = doc.add_heading('✨ 今年度の改善ポイント', level=2)
    heading3.runs[0].font.size = Pt(13)
    heading3.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading3.paragraph_format.space_after = Pt(4)
    
    improvements = [
        ('早期配布（3/1）', '総会の4週間前。家でゆっくり読めます'),
        ('質疑応接（3/7, 3/14）', '集会所で個別に相談できます'),
        ('意見用紙', '文書で自分のペースで意見を出せます'),
        ('段階的決定', '急がず、しっかり考えられます'),
    ]
    
    for title, desc in improvements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).font.bold = True
        p.add_run(f'：{desc}')
        for run in p.runs:
            run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # ========== メリハリ ==========
    heading4 = doc.add_heading('⚖️ すべてに時間をかけるわけではない', level=2)
    heading4.runs[0].font.size = Pt(12)
    heading4.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading4.paragraph_format.space_after = Pt(4)
    
    balance_table = doc.add_table(rows=5, cols=2)
    balance_table.style = 'Light List Accent 1'
    
    hdr = balance_table.rows[0].cells
    hdr[0].text = '案件の種類'
    hdr[1].text = '決め方'
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    balance_data = [
        ('日常的なこと', '役員判断'),
        ('従来通りのこと', '簡易承認'),
        ('新しい試み', '意見収集'),
        ('重要な変更', '十分な時間'),
    ]
    
    for i, (kind, method) in enumerate(balance_data, start=1):
        cells = balance_table.rows[i].cells
        cells[0].text = kind
        cells[1].text = method
        for cell in cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # ========== よくある心配 ==========
    heading5 = doc.add_heading('❓ よくある心配', level=2)
    heading5.runs[0].font.size = Pt(12)
    heading5.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading5.paragraph_format.space_after = Pt(3)
    
    qa_items = [
        ('Q: 時間がかかりすぎるのでは？', 'A: 重要な案件だけです。メリハリが大切。'),
        ('Q: 反対意見があると決まらない？', 'A: 反対意見も含めて検討し、理由を説明して決めます。'),
    ]
    
    for q, a in qa_items:
        p_q = doc.add_paragraph()
        p_q.add_run(q).font.bold = True
        p_q.runs[0].font.size = Pt(9)
        p_q.paragraph_format.space_after = Pt(1)
        
        p_a = doc.add_paragraph(a)
        p_a.runs[0].font.size = Pt(9)
        p_a.paragraph_format.space_after = Pt(3)
        p_a.paragraph_format.left_indent = Inches(0.2)
    
    # ========== まとめ ==========
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    summary_box = doc.add_paragraph()
    summary_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_box.add_run('💡 決め方を変えると、自治会が変わる')
    summary_run.font.bold = True
    summary_run.font.size = Pt(12)
    summary_run.font.color.rgb = RGBColor(0, 51, 102)
    summary_box.paragraph_format.space_after = Pt(3)
    
    summary_text = doc.add_paragraph('「どう決めるか」を一緒に考えることが、自治会を良くする第一歩です。')
    summary_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_text.runs[0].font.size = Pt(10)
    summary_text.paragraph_format.space_after = Pt(6)
    
    # ========== 意見募集 ==========
    heading6 = doc.add_heading('📝 あなたの意見をお聞かせください', level=2)
    heading6.runs[0].font.size = Pt(11)
    heading6.runs[0].font.color.rgb = RGBColor(0, 102, 51)
    heading6.paragraph_format.space_after = Pt(3)
    
    contact_items = [
        '集会所での相談：3月7日・14日',
        '意見用紙の提出：3月21日締切',
        '総会での発言：3月28日',
    ]
    
    for item in contact_items:
        p = doc.add_paragraph(f'• {item}', style='List Bullet')
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)
    
    # ========== フッター ==========
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    footer = doc.add_paragraph('2026年2月　観音台第二自治会 事務局')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # ========== 保存 ==========
    output_path = os.path.join('docs', 'community', '2026__', 'decision_making_process.docx')
    doc.save(output_path)
    print(f"\n✓ SUCCESS! Word文書を作成しました")
    print(f"✓ Output: {output_path}")
    
    return output_path

if __name__ == "__main__":
    create_decision_making_document()
