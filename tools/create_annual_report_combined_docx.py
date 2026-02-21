r"""
年報 三部構成見本 生成スクリプト

【構成】
  第1部（風土の共有と財政の見通し）  ← 実際の内容（ドラフトから読み込み）
  第2部（住民アンケート記録）        ← 実際の内容（ドラフトから読み込み）
  第3部（総会資料）                  ← 構成イメージのみ（仮データの確認チェックリスト付き）

【AI assisted 自治会運営の設計思想】
  テキストを修正したい → 各ドラフト .md を編集してこのスクリプトを再実行
  来年度に引き継ぐ場合 → 同じ手順で再生成できる

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_annual_report_combined_docx.py

データソース:
    docs/community/2026__/annual_report_part1_draft_v2.md        ← 第1部前半
    docs/community/2026__/annual_report_part1_draft_v2_second.md ← 第1部後半
    docs/community/2026__/annual_report_part2_anketo_draft.md    ← 第2部

出力:
    docs/community/2026__/annual_report_combined.docx
"""

import io
import re
from pathlib import Path

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# パス設定
# ============================================================
BASE              = Path(__file__).parent.parent
DRAFT_P1_FIRST    = BASE / 'docs/community/2026__/annual_report_part1_draft_v2.md'
DRAFT_P1_SECOND   = BASE / 'docs/community/2026__/annual_report_part1_draft_v2_second.md'
DRAFT_P2          = BASE / 'docs/community/2026__/annual_report_part2_anketo_draft.md'
OUTPUT            = BASE / 'docs/community/2026__/annual_report_combined.docx'

IMG_HIERARCHY     = BASE / 'docs/community/2026__/images/hierarchy_diagram.png'
IMG_FORECAST      = BASE / 'tools/forecast_garbage_cumulative.png'
IMG_COVER         = BASE / 'docs/community/2026__/images/report_cover.jpg'

# 色
NAVY      = RGBColor(0x17, 0x37, 0x5E)
MID_NAVY  = RGBColor(0x26, 0x4F, 0x78)
GREY      = RGBColor(0x60, 0x60, 0x60)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
AMBER     = RGBColor(0x80, 0x40, 0x00)  # 確認事項の色

# ============================================================
# 共通ユーティリティ
# ============================================================

def set_default_font(doc, font_name='游明朝', size_pt=11):
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    for tag, attr, val in [
        ('w:fldChar',   'w:fldCharType', 'begin'),
        ('w:instrText', 'xml:space',     'preserve'),
        ('w:fldChar',   'w:fldCharType', 'end'),
    ]:
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.text = 'PAGE'
        el.set(qn(attr), val)
        run._r.append(el)


def parse_bold(text):
    """**bold** → [(text, is_bold), ...]"""
    result, pos = [], 0
    for m in re.finditer(r'\*\*(.*?)\*\*', text):
        if m.start() > pos:
            result.append((text[pos:m.start()], False))
        result.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        result.append((text[pos:], False))
    return result or [(text, False)]


def add_para(doc, text, font_size=11, indent_left=0,
             space_before=None, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if indent_left:
        p.paragraph_format.left_indent = Inches(indent_left)
    for seg, bold in parse_bold(text):
        run = p.add_run(seg)
        run.bold = bold
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
    return p


def add_heading(doc, text, level, font_size, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 3 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = color or (NAVY if level <= 3 else MID_NAVY)
    return p


def add_shaded_block(doc, text, fill='F0F0F0', font_size=9, color=None):
    """グレー背景のブロックを挿入"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(16)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return p


def add_chart_placeholder(doc, key):
    LABELS = {
        '夏_空き地利用': '図：空き地利用への意向（夏アンケート）',
        '夏_地域要望':   '図：地域要望の集計（夏アンケート・複数回答可）',
        '冬_会費改定':   '図：会費改定案への意向（冬アンケート・3択）',
        '冬_参考項目':   '図：参考項目の集計（役員会委任・前提支持）',
    }
    label = LABELS.get(key, f'図：{key}')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'D9D9D9')
    pPr.append(shd)
    run = p.add_run(f'▼ {label} ▼\n（ここにExcelの図を貼り付けてください）')
    run.font.size = Pt(10)
    run.font.color.rgb = GREY


def insert_image(doc, img_path, width=5.5, caption=None):
    """画像を挿入してキャプションを付ける"""
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(8)
    ip.paragraph_format.space_after  = Pt(4)
    ip.add_run().add_picture(str(img_path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        rc = cap.add_run(caption)
        rc.font.size = Pt(9)
        rc.font.color.rgb = GREY


def render_table(doc, rows):
    """Markdownテーブル行リストをWordテーブルに変換"""
    data = []
    for row in rows:
        if re.match(r'\|[\s\-:|]+\|', row):
            continue
        cells = [c.strip() for c in row.strip('|').split('|')]
        data.append(cells)
    if not data:
        return
    ncols = max(len(r) for r in data)
    tbl = doc.add_table(rows=len(data), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            if ci >= ncols:
                break
            cell = tbl.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            for t, bold in parse_bold(cell_text):
                r = p.add_run(t)
                r.bold = bold or (ri == 0)
                r.font.size = Pt(10)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


def new_section(doc):
    """新しいセクションを追加（A4・標準余白）"""
    sec = doc.add_section()
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)
    return sec


def add_part_header_inline(doc, part_num_ja, title, subtitle=None):
    """各部のヘッダを本文先頭にインライン配置（ページブレークなし）"""
    # 第○部 ラベル（小）
    pl = doc.add_paragraph()
    pl.paragraph_format.space_before = Pt(2)
    pl.paragraph_format.space_after  = Pt(2)
    rl = pl.add_run(f'第{part_num_ja}部')
    rl.bold = True
    rl.font.size = Pt(10)
    rl.font.color.rgb = GREY

    # タイトル（大）
    pt = doc.add_paragraph()
    pt.paragraph_format.space_before = Pt(0)
    pt.paragraph_format.space_after  = Pt(4)
    rt = pt.add_run(title)
    rt.bold = True
    rt.font.size = Pt(20)
    rt.font.color.rgb = NAVY

    # サブタイトル
    if subtitle:
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after  = Pt(12)
        ps.add_run(subtitle).font.size = Pt(10)

    # 区切り線
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(0)
    p_hr.paragraph_format.space_after  = Pt(10)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '6')
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), '173758')  # NAVY
    pBdr.append(btm)
    pPr.append(pBdr)


# ============================================================
# 第1部パーサー（draft_v2.md の「## ドラフト本文」）
# ============================================================

def parse_part1_first(md_path):
    text = md_path.read_text(encoding='utf-8')
    m = re.search(
        r'## ドラフト本文.*?\n\n(.*?)(?=\n## 執筆メモ|\n## v1からの|\Z)',
        text, re.DOTALL
    )
    if not m:
        return []

    lines = m.group(1).split('\n')
    items, buf = [], []

    def flush():
        if buf:
            items.append({'type': 'para', 'content': '\n'.join(buf)})
            buf.clear()

    for line in lines:
        s = line.strip()
        if s == '---':
            flush()
        elif s.startswith('### '):
            flush()
            items.append({'type': 'h3', 'content': s[4:]})
        elif s.startswith('!['):
            flush()
            img_m = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', s)
            if img_m:
                items.append({'type': 'fig_hierarchy', 'alt': img_m.group(1), 'src': img_m.group(2)})
        elif s.startswith('**[図：'):
            flush()  # テキスト参照行はスキップ（直後の ![...] タグで処理）
        elif s.startswith('>'):
            part = s[1:].strip()
            if part:
                buf.append(part)
            elif buf:
                flush()
        elif s:
            buf.append(s)
        else:
            flush()
    flush()
    return items


# ============================================================
# 第1部後半パーサー（draft_v2_second.md の「## ドラフト本文」）
# ============================================================

def parse_part1_second(md_path):
    text = md_path.read_text(encoding='utf-8')
    m = re.search(r'## ドラフト本文.*?\n\n(.*?)(?=\n## 執筆メモ|\Z)', text, re.DOTALL)
    if not m:
        return []

    lines = m.group(1).split('\n')
    items, buf = [], []

    def flush():
        if buf:
            items.append({'type': 'para', 'content': '\n'.join(buf)})
            buf.clear()

    for line in lines:
        s = line.strip()
        if s == '---':
            flush()
        elif s.startswith('### '):
            flush()
            items.append({'type': 'h3', 'content': s[4:]})
        elif s.startswith('|') and len(s) > 1:
            flush()
            if items and items[-1]['type'] == 'table':
                items[-1]['rows'].append(s)
            else:
                items.append({'type': 'table', 'rows': [s]})
        elif s.startswith('!['):
            flush()
            items.append({'type': 'fig_forecast'})
        elif s.startswith('**[図：'):
            flush()  # スキップ（直後の ![...] で処理）
        elif s:
            buf.append(s)
        else:
            flush()
    flush()
    return items


# ============================================================
# 第2部パーサー（annual_report_part2_anketo_draft.md）
# ============================================================

def parse_part2(md_path):
    text = md_path.read_text(encoding='utf-8')
    m = re.search(r'## ドラフト本文\n+(.*?)(?=\n## 執筆メモ|\Z)', text, re.DOTALL)
    if not m:
        raise ValueError('ドラフト本文セクションが見つかりません')

    lines = m.group(1).split('\n')
    items, buf = [], []

    def flush():
        if buf:
            items.append({'type': 'para', 'content': '\n'.join(buf)})
            buf.clear()

    for line in lines:
        s = line.strip()
        if s == '---':
            flush()
        elif s.startswith('#### '):
            flush()
            items.append({'type': 'h4', 'content': s[5:]})
        elif s.startswith('### '):
            flush()
            items.append({'type': 'h3', 'content': s[4:]})
        elif s.startswith('## '):
            flush()
            items.append({'type': 'h2', 'content': s[3:]})
        elif s.startswith('[CHART:') and s.endswith(']'):
            flush()
            items.append({'type': 'chart', 'key': s[7:-1]})
        elif s.startswith('>'):
            part = s[1:].strip()
            if part:
                buf.append(part)
            elif buf:
                flush()
        elif s:
            buf.append(s)
        else:
            flush()
    flush()
    return items


# ============================================================
# 第1部免責文（draft_v2.md の「編集者より」ブロック）
# ============================================================

def extract_disclaimer(md_path):
    text = md_path.read_text(encoding='utf-8')
    m = re.search(r'> \*\*編集者より\*\*\s*\n((?:>.*\n)*)', text)
    if not m:
        return None
    lines = []
    for line in m.group(0).split('\n'):
        s = line.strip()
        if s.startswith('>'):
            part = s[1:].strip()
            if part and '**編集者より**' not in part:
                lines.append(part)
    return '\n'.join(lines) if lines else None


# ============================================================
# 表紙・キャッチページ・裏表紙（PIL 画像合成）
# ============================================================

STEEL  = RGBColor(0x1A, 0x5E, 0x9B)   # キャッチ文字色（スチールブルー）
STEEL2 = RGBColor(0x70, 0xA0, 0xC8)   # "with" の薄め色

# PIL 画像入力解像度（DPI）: A4 at 150 DPI = 1240×1754 px
_COVER_DPI = 150
_COVER_W   = 1240
_COVER_H   = 1754


def _find_jp_font(px_size):
    """利用可能な Windows 日本語フォントを探して返す"""
    if not PIL_AVAILABLE:
        return None
    candidates = [
        'C:/Windows/Fonts/meiryo.ttc',
        'C:/Windows/Fonts/YuGothM.ttc',
        'C:/Windows/Fonts/msgothic.ttc',
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, px_size)
        except Exception:
            pass
    return ImageFont.load_default()


def _compose_cover(text_lines, bg_path=None, overlay_alpha=170):
    """
    PIL でカバー画像を合成して BytesIO で返す。

    text_lines: [(text, y_frac, pt_size, (R,G,B), italic), ...]
      y_frac : ページ高さに対する相対位置 (0.0−1.0)
      pt_size: ポイントサイズ（自動でピクセル数に変換）
      (R,G,B): 色
      italic : 未使用（PIL イタリック非対応のため）
    """
    if not PIL_AVAILABLE:
        return None

    W, H = _COVER_W, _COVER_H

    # 背景画像
    if bg_path and Path(bg_path).exists():
        bg = Image.open(bg_path).convert('RGB').resize((W, H), Image.LANCZOS)
        # 暗色半透明オーバーレイ（文字読みやすく）
        ov = Image.new('RGBA', (W, H), (10, 20, 50, overlay_alpha))
        bg = Image.alpha_composite(bg.convert('RGBA'), ov).convert('RGB')
    else:
        bg = Image.new('RGB', (W, H), (23, 55, 94))  # NAVY単色

    draw = ImageDraw.Draw(bg)

    for text, y_frac, pt_size, rgb, _italic in text_lines:
        px = max(12, int(pt_size * _COVER_DPI / 72))
        font = _find_jp_font(px)
        try:
            bbox = draw.textbbox((0, 0), text, font=font)
            tw = bbox[2] - bbox[0]
        except Exception:
            tw = len(text) * px // 2
        x = max(0, (W - tw) // 2)
        y = int(H * y_frac)
        draw.text((x, y), text, fill=rgb, font=font)

    buf = io.BytesIO()
    bg.save(buf, format='JPEG', quality=92)
    buf.seek(0)
    return buf


def _add_full_page_image(doc, buf):
    """フルページインライン画像を追加（余白ゼロセクションで使用）"""
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(0)
    ip.paragraph_format.space_after  = Pt(0)
    ip.add_run().add_picture(buf, width=Inches(8.27), height=Inches(11.69))


def add_front_cover(doc):
    """表紙（背景画像 + テキスト合成）— ページブレークなし"""
    buf = _compose_cover(
        bg_path=IMG_COVER if PIL_AVAILABLE else None,
        overlay_alpha=160,
        text_lines=[
            # text,                     y_frac, pt,  (R,G,B),          italic
            ('観音台一丁目第2自治会',         0.06,  14, (200,210,230), False),
            ('年　報　2025年度',         0.11,  28, (240,245,255), False),
            ('Looking Forward',          0.24,  22, ( 80,160,210), True),
            ('with',                     0.31,  13, (110,180,230), True),
            ('Shared Diversity',         0.36,  22, ( 80,160,210), True),
            ('第１部　風土の共有と財政の見通し', 0.52,  12, (210,220,240), False),
            ('第２部　住民アンケート記録',     0.58,  12, (210,220,240), False),
            ('第３部　総会資料',             0.64,  12, (210,220,240), False),
            ('2026年3月',                  0.80,  11, (170,185,210), False),
            ('観音台一丁目第２自治会　事務局・高橋正剛', 0.84, 10, (150,170,200), False),
        ],
    )
    if buf:
        _add_full_page_image(doc, buf)
    else:
        # PIL 未利用時のフォールバック: テキストのみ表紙
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(80)
        p.add_run('年　報　2025年度').font.size = Pt(28)
    # ページブレークなし（呼び元が doc.add_section() でページを切る）


def add_catch_page(doc):
    """キャッチページ（表紙の次ページ）— ページブレークあり"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    p.paragraph_format.space_after  = Pt(16)
    r = p.add_run('Looking Forward')
    r.bold = True; r.italic = True
    r.font.size = Pt(40)
    r.font.color.rgb = STEEL

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run('with')
    r2.italic = True
    r2.font.size = Pt(20)
    r2.font.color.rgb = STEEL2

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(60)
    r3 = p3.add_run('Shared Diversity')
    r3.bold = True; r3.italic = True
    r3.font.size = Pt(40)
    r3.font.color.rgb = STEEL

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(4)
    r4 = p4.add_run('多様な参加のかたちが、自治会を前へ連れていく。')
    r4.font.size = Pt(12)
    r4.font.color.rgb = GREY

    doc.add_page_break()


def add_back_cover(doc):
    """裏表紙（背景画像 + テキスト合成）— ページブレークなし"""
    buf = _compose_cover(
        bg_path=IMG_COVER if PIL_AVAILABLE else None,
        overlay_alpha=180,
        text_lines=[
            ('観音台一丁目第2自治会',   0.22, 14, (200,210,230), False),
            ('年　報　2025年度',   0.27, 22, (240,245,255), False),
            ('Looking Forward',      0.40, 18, ( 80,160,210), True),
            ('with',                 0.46, 12, (110,180,230), True),
            ('Shared Diversity',     0.50, 18, ( 80,160,210), True),
            ('https://freesemt.github.io/kannondai-community/', 0.70,  9, (150,170,200), False),
            ('© 2026 観音台一丁目第２自治会 事務局',   0.75,  9, (150,170,200), False),
        ],
    )
    if buf:
        _add_full_page_image(doc, buf)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(120)
        p.add_run('年　報　2025年度').font.size = Pt(22)


# ============================================================
# 文書組み立て
# ============================================================

def build_combined_document():
    doc = Document()
    set_default_font(doc)

    # ── Section 0: 表紙（余白ゼロ・フッターなし） ──
    sec0 = doc.sections[0]
    sec0.page_width    = Inches(8.27)
    sec0.page_height   = Inches(11.69)
    sec0.top_margin    = Inches(0)
    sec0.bottom_margin = Inches(0)
    sec0.left_margin   = Inches(0)
    sec0.right_margin  = Inches(0)
    sec0.footer.is_linked_to_previous = False
    sec0.footer.paragraphs[0].clear()

    add_front_cover(doc)  # ページブレークなし

    # ── Section 1: 本文（標準余白・フッターにページ番号） ──
    sec2 = doc.add_section()
    sec2.page_width    = Inches(8.27)
    sec2.page_height   = Inches(11.69)
    sec2.top_margin    = Inches(1.0)
    sec2.bottom_margin = Inches(1.0)
    sec2.left_margin   = Inches(1.1)
    sec2.right_margin  = Inches(1.1)
    sec2.footer.is_linked_to_previous = False
    footer_para = sec2.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_para)

    # ============================================================
    # 第1部 インラインヘッダー + 本文
    # ============================================================
    add_part_header_inline(doc, '１', '風土の共有と財政の見通し',
                           '自治会のこれまでとこれから')

    # 免責注記
    disclaimer = extract_disclaimer(DRAFT_P1_FIRST)
    if disclaimer:
        add_shaded_block(doc, disclaimer, fill='F0F0F0', font_size=9, color=DARK_GREY)

    # 第1部 前半本文
    items_p1a = parse_part1_first(DRAFT_P1_FIRST)
    print(f'  第1部前半: {len(items_p1a)} 要素', flush=True)

    for item in items_p1a:
        t = item['type']
        if t == 'h3':
            add_heading(doc, item['content'], 3, 13)
        elif t == 'fig_hierarchy':
            if IMG_HIERARCHY.exists():
                insert_image(doc, IMG_HIERARCHY, width=5.5,
                             caption='図：自治会への期待の変化（三層モデル）')
            else:
                add_para(doc, '〔図：自治会への期待の変化（三層モデル）〕',
                         space_after=10, color=GREY)
        elif t == 'para':
            for line in item['content'].split('\n'):
                if line.strip():
                    add_para(doc, line, space_after=8)

    # 第1部 後半本文（ページ区切り後）
    doc.add_page_break()
    items_p1b = parse_part1_second(DRAFT_P1_SECOND)
    print(f'  第1部後半: {len(items_p1b)} 要素', flush=True)

    for item in items_p1b:
        t = item['type']
        if t == 'h3':
            add_heading(doc, item['content'], 3, 13)
        elif t == 'para':
            for line in item['content'].split('\n'):
                if line.strip():
                    add_para(doc, line, space_after=8)
        elif t == 'table':
            render_table(doc, item['rows'])
        elif t == 'fig_forecast':
            if IMG_FORECAST.exists():
                insert_image(doc, IMG_FORECAST, width=5.8,
                             caption='図：ゴミ集積所使用料 累計推移（上：総資産推移 ／ 下：非会員数・使用料累計）')
            else:
                add_para(doc, '〔図：ゴミ集積所使用料累計推移グラフ〕',
                         space_after=10, color=GREY)

    # ============================================================
    # 第2部 インラインヘッダー + 本文
    # ============================================================
    new_section(doc)   # ページ切り（フッターは前セクションのものを継承）
    add_part_header_inline(doc, '２', '住民アンケート記録',
                           '夏アンケート・冬アンケートの実施結果と対応経遠')

    # 第2部 本文
    items_p2 = parse_part2(DRAFT_P2)
    print(f'  第2部: {len(items_p2)} 要素', flush=True)

    for item in items_p2:
        t = item['type']
        if t == 'h2':
            add_heading(doc, item['content'], 2, 14)
        elif t == 'h3':
            add_heading(doc, item['content'], 3, 13)
        elif t == 'h4':
            add_heading(doc, item['content'], 4, 12)
        elif t == 'chart':
            add_chart_placeholder(doc, item['key'])
        elif t == 'para':
            for line in item['content'].split('\n'):
                line = line.strip()
                if line.startswith('・'):
                    add_para(doc, line, indent_left=0.2, space_after=3)
                elif line:
                    add_para(doc, line, space_after=6)

    # ============================================================
    # 第3部 インラインヘッダー + 本文
    # ============================================================
    new_section(doc)
    add_part_header_inline(doc, '３', '総会資料',
                           '2025年度事業・会計報告 ／ 2026年度役員・計画・予算案')

    # 注記
    add_para(doc,
             '【注】この第3部は構成イメージです。'
             '会計係・会長の確認を経て実際の数値・氏名に差し替えます。',
             font_size=10, color=AMBER, space_after=20)

    # 議案リスト
    agendas = [
        ('第1号議案', '2025年度 事業報告',
         '会員の交流促進 ／ 役員会開催状況（全5回）／ 住宅環境整備 ／ '
         'つくば市への要望・協力 ／ 前年度総会課題への対応'),
        ('第2号議案', '2025年度 会計報告',
         '一般会計（収入・支出・次年度繰越）／ 集会所積立金特別会計 ／ '
         '会計監査報告 ／ 募金活動まとめ'),
        ('第3号議案', '2026年度 役員改選案',
         '会長・事務局長・会計係・各班班長（6名）・会計監事（2名）の改選'),
        ('第4号議案', '2026年度 事業計画案',
         '地域環境整備 ／ 会員の交流促進 ／ '
         '役員選出の安定化・多様な参加形態の検討継続 ／ 集会所積立金の適正水準検討'),
        ('第5号議案', '2026年度 予算案',
         '収入（会費・つくば市委託金・繰越金）／ 支出（役員報酬・事業費・管理費等）'),
        ('第6号議案', '会則の一部改正',
         '（改正案が確定した場合に掲載）'),
    ]

    for num, title, desc in agendas:
        p_num = doc.add_paragraph()
        p_num.paragraph_format.space_before = Pt(14)
        p_num.paragraph_format.space_after  = Pt(2)
        rn = p_num.add_run(num)
        rn.bold = True
        rn.font.size = Pt(10)
        rn.font.color.rgb = GREY

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after  = Pt(2)
        rt = p_title.add_run(title)
        rt.bold = True
        rt.font.size = Pt(13)
        rt.font.color.rgb = NAVY

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(0)
        p_desc.paragraph_format.space_after  = Pt(2)
        p_desc.paragraph_format.left_indent  = Inches(0.2)
        rd = p_desc.add_run(desc)
        rd.font.size = Pt(10)
        rd.font.color.rgb = DARK_GREY

        # 区切り線
        p_hr = doc.add_paragraph()
        p_hr.paragraph_format.space_after = Pt(0)
        pPr = p_hr._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # 確認が必要な仮データ一覧
    add_para(doc, '', space_before=24, space_after=4)  # 空白
    add_shaded_block(doc,
        '【仮データ確認チェックリスト】\n'
        '□ 入退会者リスト（氏名・班・日付・区分）\n'
        '□ 役員会各回の正確な日付（第1回・第3回）\n'
        '□ 住宅環境整備の詳細（ゴミ集積所清掃用具・出光交差点安全対策の顛末）\n'
        '□ 第一自治会との合同役員会（実施/未実施・内容）\n'
        '□ つくば市への要望・協力の具体的案件\n'
        '□ 会計報告の実数値（会計係から受領後に反映）\n'
        '□ 役員改選の候補者氏名（石川さん・加藤さん・次期班長6名）\n'
        '□ 会計監査報告の日付・監事氏名',
        fill='FFF3E0', font_size=10, color=AMBER)

    # ── Section last: 裏表紙（余白ゼロ・フッターなし） ──
    sec_back = doc.add_section()
    sec_back.page_width    = Inches(8.27)
    sec_back.page_height   = Inches(11.69)
    sec_back.top_margin    = Inches(0)
    sec_back.bottom_margin = Inches(0)
    sec_back.left_margin   = Inches(0)
    sec_back.right_margin  = Inches(0)
    sec_back.footer.is_linked_to_previous = False
    sec_back.footer.paragraphs[0].clear()

    add_back_cover(doc)  # ページブレークなし

    return doc


# ============================================================
# メイン処理
# ============================================================

def main():
    print('=== 年報 三部構成見本 生成スクリプト ===', flush=True)
    print(f'第1部前半: {DRAFT_P1_FIRST.name}', flush=True)
    print(f'第1部後半: {DRAFT_P1_SECOND.name}', flush=True)
    print(f'第2部    : {DRAFT_P2.name}', flush=True)

    doc = build_combined_document()

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f'[OK] 出力完了: {OUTPUT}', flush=True)


if __name__ == '__main__':
    main()
