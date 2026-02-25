r"""
角倉会長への連絡書 生成スクリプト

Usage:
    cd E:\GitHub\kannondai-community
    & "C:\Program Files\Python313\python.exe" tools\create_chairman_letter_docx.py

出力:
    docs/community/2026__/chairman_letter_20260225.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).parent.parent / 'docs/community/2026__/chairman_letter_20260225.docx'

DARK = RGBColor(0x20, 0x20, 0x20)


def set_default_font(doc, font_name='游明朝', size_pt=11):
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def set_run_font(run, font_name='游明朝'):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def para(doc, text='', size=11, bold=False, align=None,
         indent=0, space_before=0, space_after=6, color=None):
    """テキストを段落として追加。**...** をインライン太字に処理。"""
    import re
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    segments = re.split(r'(\*\*.*?\*\*)', text)
    for seg in segments:
        m = re.match(r'\*\*(.*?)\*\*', seg)
        content = m.group(1) if m else seg
        run = p.add_run(content)
        run.bold = bool(m) or bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        set_run_font(run)
    return p


def hline(doc):
    """水平罫線（段落の下罫線で代用）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def build():
    doc = Document()

    # ページ余白（A4）
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    set_default_font(doc)

    # ---- タイトル ----
    para(doc, '連　絡　書', size=16, bold=True, align='center',
         space_before=0, space_after=16)

    # ---- 日付・差出人（右寄せ） ----
    para(doc, '令和8年（2026年）2月25日', size=11, align='right', space_after=2)
    para(doc, '会長　角倉雅博　様', size=11, align='right', space_after=2)
    para(doc, '事務局長　高橋 正剛', size=11, align='right', space_after=12)

    hline(doc)

    # ---- 書き出し ----
    para(doc,
         '先日は「年報 2025（案）」草稿へのご意見をいただき、ありがとうございました。',
         size=11, space_before=10, space_after=8)

    para(doc,
         '**「役員間で合意されていない内容を今年度総会資料とするのは妥当ではない」**'
         'というご指摘は、ごもっともと受け止めております。',
         size=11, space_after=8)

    para(doc, 'ご返信を踏まえ、方針を以下のとおりにいたします。',
         size=11, space_after=10)

    hline(doc)

    # ---- 1項 ----
    para(doc, '**1．3月28日（土）の総会に提出する議案**',
         size=11, space_before=10, space_after=4)
    para(doc,
         '　総会資料は従来通りの構成（事業報告・決算・役員選任など）とします。'
         '年報の内容は、総会議案には含めません。',
         size=11, indent=0.5, space_after=10)

    # ---- 2項 ----
    para(doc, '**2．年報について**',
         size=11, space_before=4, space_after=4)
    para(doc,
         '　「年報 2025 案」として、役員会の承認を経ない**参考資料**として配布します。'
         '役員会としての決定事項ではなく、事務局による問題提起・記録としての位置づけです。'
         '承認をお願いするものではありませんので、会長にご負担をおかけするものではありません。',
         size=11, indent=0.5, space_after=6)
    para(doc,
         '　この資料を作成したのは、今年度は会費見直しや役員選出の困難など、'
         '例年の総会資料では説明しきれない事情が重なっているためです。'
         '会員の方々に状況を正確に知っていただきたいという趣旨で、記録として残したいと考えました。',
         size=11, indent=0.5, space_after=10)

    # ---- 3項 ----
    para(doc, '**3．次年度役員候補との事前相談**',
         size=11, space_before=4, space_after=4)
    para(doc,
         '　3月1日（日）に、来年度の役員候補（中山氏・石川氏・加藤氏）と'
         '事前に意見交換の場を設ける予定です。',
         size=11, indent=0.5, space_after=10)

    hline(doc)

    # ---- 結び ----
    para(doc, 'ご返信は不要です。ご承知おきいただければ幸いです。',
         size=11, space_before=10, space_after=0)

    doc.save(OUTPUT)
    print(f'✅ 保存しました: {OUTPUT}')


if __name__ == '__main__':
    build()
